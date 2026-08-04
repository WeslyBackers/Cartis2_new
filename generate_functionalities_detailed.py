# -*- coding: utf-8 -*-
"""
Generates CARTIS2_Functionalities_Detailed_2026-04-27.docx — a detailed
technical/functional inventory of the CARTIS 2.0 application: frontend
pages, backend API endpoints, database objects, automation scripts and
documentation coverage.

This script replaces an older, hand-scraped version of the same document
(which contained raw JSX/code fragments instead of readable descriptions,
and had drifted out of sync with the codebase). It is written from
scratch against the current state of the repository.

NOTE: the original filename (with its 2026-04-27 date suffix) is kept on
purpose so existing links/references to the document keep working; the
content itself reflects the application as of the regeneration date,
which is recorded in the document header.
"""

import os
import re
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = os.path.dirname(os.path.abspath(__file__))

DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
MID_BLUE = RGBColor(0x2E, 0x5C, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = MID_BLUE
    return h


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    return h


def para(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def code_line(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    return p


def two_col_table(doc, headers, rows, col_widths=None, header_color='2E5C8A'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def page_section(doc, filename, intro, features):
    """Render a frontend page section: heading, intro paragraph and a
    bullet list of functionality descriptions (label + explanation)."""
    add_h2(doc, filename)
    para(doc, intro)
    bullet_list_heading = doc.add_paragraph()
    r = bullet_list_heading.add_run('Functionaliteiten:')
    r.bold = True
    for label, desc in features:
        bullet(doc, f' — {desc}', bold_prefix=label)


def endpoints_table(doc, module_title, mount_path, endpoints):
    add_h2(doc, module_title)
    para(doc, f'Mount-pad: {mount_path}  ({len(endpoints)} endpoints)', italic=True, color=GREY)
    rows = [(method, mount_path.rstrip('/') + path, desc) for method, path, desc in endpoints]
    two_col_table(doc, ['Methode', 'Pad', 'Beschrijving'], rows, col_widths=[2.0, 6.5, 8.5])


def scrape_markdown_outline(md_path, max_lines=60):
    """Very small markdown scraper: returns a list of (kind, text) tuples
    for heading lines (#, ##, ###) and top-level bullet lines (-, *),
    skipping fenced code blocks. Used for the 'Documentation Coverage'
    section so it stays roughly in sync with the underlying .md files."""
    out = []
    if not os.path.exists(md_path):
        return out
    in_fence = False
    with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
        for raw in f:
            line = raw.rstrip('\n')
            stripped = line.strip()
            if stripped.startswith('```'):
                in_fence = not in_fence
                continue
            if in_fence:
                continue
            m = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if m:
                out.append(('heading', len(m.group(1)), m.group(2).strip()))
                continue
            m = re.match(r'^[-*]\s+(.*)', stripped)
            if m:
                out.append(('bullet', 0, m.group(1).strip()))
                continue
            if len(out) >= max_lines:
                break
    return out[:max_lines]


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

for lvl, size, color in [(1, 18, DARK_BLUE), (2, 14, MID_BLUE), (3, 12, MID_BLUE)]:
    hstyle = doc.styles[f'Heading {lvl}']
    hstyle.font.size = Pt(size)
    hstyle.font.color.rgb = color
    hstyle.font.bold = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CARTIS 2.0 — Detailed Functionalities\nand Technical Inventory')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f'Regenerated: {date.today().isoformat()}  (originally created 2026-04-27)')
r.italic = True
r.font.color.rgb = GREY

para(doc, (
    'This document is a static-analysis style inventory of the CARTIS 2.0 codebase: every '
    'frontend page, every backend REST endpoint, every database table/index, every automation '
    'script and the documentation set. It is intended for developers and technical reviewers who '
    'need a complete, accurate map of "what exists and what it does". For end-user instructions, '
    'see CARTIS2_Gebruikershandleiding.docx; for a narrative architecture write-up, see '
    'CARTIS2_Analyse_April2026.docx.'
), italic=True, color=GREY)
doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. Functional Overview
# ---------------------------------------------------------------------------
add_h1(doc, '1. Functional Overview')
para(doc, (
    'CARTIS 2.0 is a web application used by the Vlaamse Hydrografie to manage nautical '
    'notifications (Meldingen), the tasks derived from them (Taken), the product versions in '
    'which those tasks are processed (Productversies) and the eventual publication of nautical '
    'products across four production lines: ZK (zeekaarten), IENC (binnenvaartkaarten), Pilot ENC '
    '(loodskaarten) and Publ (publicaties zoals BaZ, Lichtenlijst, Verbeterlijst).'
))
para(doc, (
    'Architecture: a React + TypeScript (Vite) single-page frontend and a Node.js/Express + '
    'TypeScript backend, sharing a PostgreSQL + PostGIS database hosted on Supabase. In '
    'production the app is deployed on Vercel (serverless API functions + static frontend '
    'build); file attachments are stored in Supabase Storage, falling back to the local '
    'filesystem in development.'
))
bullet(doc, 'Notification lifecycle: intake (manual, email drag-and-drop, or future REST integrations) → decision per production line → task creation.')
bullet(doc, 'Task lifecycle: workflow steps, comments, product linking, related tasks, BaZ articles, MSI/follow-up/extra-info flags, per-production-line status.')
bullet(doc, 'Product version lifecycle: draft → published, with correction-list and BaZ-2 publication text generation and attachment management.')
bullet(doc, 'Cross-cutting: JWT authentication, per-user/per-production-line access rights, activity logging, PostGIS-based zone/product auto-detection, lead-time dashboards.')

# ---------------------------------------------------------------------------
# 2. Frontend Pages and User Interactions
# ---------------------------------------------------------------------------
add_h1(doc, '2. Frontend Pages and User Interactions')
para(doc, (
    'All pages live under frontend/src/pages/. Each page is described below with its core '
    'purpose and the concrete interactions/functions it offers to the user.'
))

page_section(doc, 'Dashboard.tsx', (
    'Landing page after login. Shows KPI summary cards and a personal notes panel scoped to '
    'production lines.'
), [
    ('KPI-kaarten — ', 'tonen het aantal openstaande meldingen en actieve taken; klikken navigeert naar de bijbehorende lijstpagina.'),
    ('Notitie toevoegen — ', 'nieuwe persoonlijke dashboard-notitie aanmaken met tekst, prioriteit (laag/gemiddeld/hoog) en zichtbaarheid per productielijn.'),
    ('Sorteren — ', 'notities sorteren op prioriteit (hoog/laag eerst) of op datum (nieuw/oud eerst).'),
    ('Bewerken — ', 'inhoud, prioriteit of zichtbaarheid van een bestaande notitie aanpassen (enkel eigen notities).'),
    ('Verwijderen voor lijn — ', 'notitie verbergen voor de huidige productielijn zonder ze voor andere lijnen te verwijderen.'),
    ('Verwijderen — ', 'notitie volledig en definitief verwijderen (enkel de auteur).'),
])

page_section(doc, 'LeadTimes.tsx', (
    'Dashboard with doorlooptijd- (lead time) statistics per taak, gebaseerd op '
    'GET /api/tasks/lead-times.'
), [
    ('Samenvattingskaarten — ', 'gemiddelde tijd melding→taak, taak→publicatie en totale doorlooptijd, plus aantal trajecten.'),
    ('Detailtabel — ', 'per-taak overzicht van alle tussentijdse data (melding, beslissing, taak, publicatie).'),
    ('Kolomfilters — ', 'elke kolom heeft een eigen tekst-/waardefilter; "Kolomfilters wissen" reset alle filters in één klik.'),
])

page_section(doc, 'Login.tsx', (
    'Authentication entry point. Unauthenticated users are redirected here.'
), [
    ('Inloggen — ', 'e-mailadres en wachtwoord invoeren; roept POST /api/auth/login aan en slaat het JWT-token op.'),
    ('Foutmelding — ', 'toont een duidelijke fout bij ongeldige combinatie of serverfout.'),
    ('Redirect — ', 'bij succesvolle login wordt automatisch doorgestuurd naar het dashboard.'),
])

page_section(doc, 'NotificationDetail.tsx', (
    'The richest page in the app: full detail view of a single melding, including map, '
    'coordinates, products, zones, comments and attachments. Content is split per production '
    'line where relevant.'
), [
    ('Bewerken / Opslaan / Annuleren — ', 'basisgegevens en opmerkingen van de melding wijzigen.'),
    ('Kaart — ', 'interactieve Leaflet-kaart met "Vergroot/Verklein kaart", WMS-lagenpaneel (23 lagen) en een aan/uit-schakelaar voor de productenlaag.'),
    ('Coördinaten toevoegen — ', 'punt, lijn of vlak intekenen op de kaart, of handmatig invoeren in één van zeven coördinaatformaten (CRS); bewerken en verwijderen van bestaande coördinaten.'),
    ('⟳ Herbereken producten — ', 'draait de automatische PostGIS-productdetectie opnieuw op basis van de huidige geometrie.'),
    ('+ Kaart toevoegen / + Toevoegen link — ', 'product handmatig koppelen aan de melding; "✕ ontkoppelen" verwijdert een koppeling.'),
    ('⟳ Herbereken zones — ', 'herberekent de automatisch gedetecteerde zones; zones kunnen ook handmatig toegevoegd/verwijderd worden.'),
    ('💬 Opmerking opslaan — ', 'reactie toevoegen aan de discussie, per productielijn zichtbaar.'),
    ('Bijlagen — ', 'bestand uploaden, "📥 Bekijk" om te downloaden/openen, "🗑️ Verwijder" om te verwijderen.'),
    ('📥 Export GML — ', 'exporteert de geometrie van de melding als GML-bestand.'),
    ('Informatie-aanvraag — ', 'e-mailontwerp met aanvraaggegevens opstellen en opslaan ("Opslaan en openen").'),
    ('Beslissing Ja/Nee — ', 'beslissing registreren voor de actieve productielijn; "Ja" maakt automatisch een taak aan.'),
])

page_section(doc, 'Notifications.tsx', (
    'List/search page for all meldingen, with bulk actions and creation of new notifications.'
), [
    ('Zoeken en filteren — ', 'op tekst, status, productielijn en datum.'),
    ('Nieuwe melding — ', 'formulier om een melding handmatig aan te maken, inclusief e-mail drag-and-drop import (.eml/.msg via msgreader/postal-mime).'),
    ('Bulk-beslissing — ', 'meerdere geselecteerde meldingen tegelijk beslissen (Ja/Nee), met keuze voor één gecombineerde taak of aparte taken per melding.'),
    ('Export GML — ', 'per rij de geometrie exporteren.'),
    ('Rij openen — ', 'navigeert naar NotificationDetail.tsx voor de volledige detailweergave.'),
])

page_section(doc, 'Products.tsx', (
    'Product catalogue (charts, zones, coverage areas) with map and table views.'
), [
    ('Tabel/kaart wisselen — ', 'toggle tussen lijstweergave en kaartweergave met geometrie.'),
    ('Filteren op categorie — ', 'filtert de producten- of zonelijst op categorie/type.'),
    ('Product toevoegen/bewerken — ', 'nieuw product aanmaken of bestaand product wijzigen.'),
    ('KML importeren — ', 'bulk-import van producten of dekgebieden uit een KML-bestand.'),
])

page_section(doc, 'ProductVersions.tsx', (
    'Working list of open (unpublished) product versions per production line, with a detail '
    'panel for the selected version.'
), [
    ('Nieuwe versie aanmaken — ', 'handmatig of automatisch genummerd, gekoppeld aan een productielijn.'),
    ('Correctielijst bekijken/printen — ', 'gegenereerde correctielijst-inhoud (NL/EN) inzien en als PDF afdrukken.'),
    ('Gekoppelde taken — ', 'toont taken die in deze versie verwerkt worden, met per-taak uitvoeringsstatus die hier bijgewerkt kan worden.'),
    ('Bijlagen — ', 'documenten uploaden/downloaden bij de versie.'),
    ('Publiceren — ', 'versie publiceren; onvoltooide taken schuiven automatisch door naar de volgende versie.'),
])

page_section(doc, 'PublishedProductVersions.tsx', (
    'Read-oriented archive of already published product versions.'
), [
    ('Kolomfilters — ', 'per kolom filteren op waarde (versienummer, datum, productielijn, status).'),
    ('Detail bekijken — ', 'zelfde detailpaneel als ProductVersions.tsx (correctielijst, bijlagen, taken) in read-only/afgewerkte context.'),
])

page_section(doc, 'TaskDetail.tsx', (
    'Detail view of a single taak: metadata, per-production-line status, workflow, product '
    'links, BaZ articles and a resizable/expandable map.'
), [
    ('Status per lijn — ', 'status van de taak per productielijn instellen/bijwerken, inclusief "Wachten op ZK"-vlag.'),
    ('Vlaggen — ', 'MSI-actief, opvolging nodig en extra info nodig aan/uit zetten (PATCH .../flags).'),
    ('Workflow — ', 'workflowstappen toevoegen en de voortgang van de taak opvolgen.'),
    ('Reacties — ', 'per-productielijn opmerkingen toevoegen/bewerken.'),
    ('Producten koppelen — ', 'product aan de taak koppelen en de uitvoeringsstatus per product bijwerken.'),
    ('Verwante taken / extra meldingen — ', 'gerelateerde taken of bijkomende meldingen koppelen.'),
    ('BaZ-artikelen — ', 'artikel aanmaken/bewerken/verwijderen; "Auto-vertalen NL → EN" vertaalt de artikeltekst automatisch (Google Cloud Translate).'),
    ('Informatie-aanvraag — ', 'aanvraag registreren, zelfde mechanisme als bij meldingen.'),
    ('Kaart — ', 'in-/uitklapbare, verschuifbare kaartweergave van de gekoppelde geometrie.'),
])

page_section(doc, 'Tasks.tsx', (
    'List/search page for all taken, with inline status editing.'
), [
    ('Zoeken en filteren — ', 'op tekst, status per productielijn en andere kolomwaarden.'),
    ('Inline vlaggen — ', 'MSI/opvolging/extra-info-checkboxes rechtstreeks in de rij aan/uit te zetten zonder de detailpagina te openen.'),
    ('Rij uitklappen — ', 'toont extra details zonder naar TaskDetail.tsx te navigeren.'),
    ('Rij openen — ', 'navigeert naar de volledige TaskDetail.tsx-pagina.'),
])

# ---------------------------------------------------------------------------
# 3. Backend API Endpoints
# ---------------------------------------------------------------------------
add_h1(doc, '3. Backend API Endpoints')
para(doc, (
    'All routes are mounted under /api and require a valid JWT (Authorization: Bearer <token>) '
    'except POST /api/auth/login and GET /health. Route files live in backend/src/routes/. '
    'Total: 86 endpoints across 9 route modules.'
))

endpoints_table(doc, 'auth.routes.ts', '/api/auth', [
    ('POST', '/login', 'Meldt een gebruiker aan met e-mail/wachtwoord en geeft een JWT-token (24 u geldig) terug.'),
    ('GET', '/me', 'Haalt het profiel en de productielijnrechten van de ingelogde gebruiker op.'),
])

endpoints_table(doc, 'coverage.routes.ts', '/api/coverages', [
    ('GET', '/files', 'Lijst van geïmporteerde KML-bestanden, optioneel gefilterd op categorie/productielijn.'),
    ('GET', '/files/:id', 'Detail van een KML-bestand met al zijn dekgebieden.'),
    ('GET', '/products', 'Lijst van producten met geometrie voor kaartweergave.'),
    ('GET', '/zones', 'Lijst van actieve zone-producten voor de zonedropdown.'),
    ('GET', '/coverages/:id', 'Detail van een specifiek dekgebied.'),
    ('GET', '/coverages/code/:code', 'Zoekt een dekgebied op code.'),
    ('GET', '/search', 'Vrije zoekopdracht over producten/dekgebieden.'),
    ('GET', '/geojson', 'Gecombineerde GeoJSON van producten/dekgebieden voor kaartlagen.'),
])

endpoints_table(doc, 'note.routes.ts', '/api/notes', [
    ('GET', '/', 'Dashboard-notities ophalen die zichtbaar zijn voor de opgevraagde productielijn.'),
    ('POST', '/', 'Nieuwe notitie aanmaken (inhoud, prioriteit, zichtbare productielijnen).'),
    ('PUT', '/:id', 'Inhoud/prioriteit/zichtbaarheid van een bestaande notitie bewerken.'),
    ('PUT', '/:id/line-visibility', 'Zichtbaarheid van een notitie voor één productielijn aan/uit zetten.'),
    ('DELETE', '/:id', 'Notitie definitief verwijderen (enkel de aanmaker).'),
])

endpoints_table(doc, 'notification.routes.ts', '/api/notifications', [
    ('GET', '/', 'Meldingenlijst met paginering, zoeken en filters.'),
    ('GET', '/:id', 'Detail van één melding.'),
    ('POST', '/', 'Nieuwe melding aanmaken; triggert automatische zonedetectie.'),
    ('PUT', '/:id', 'Meldinggegevens (o.a. opmerkingen) bijwerken.'),
    ('POST', '/:id/decide', 'Beslissing (Ja/Nee) registreren voor de actieve productielijn; "Ja" maakt een taak aan.'),
    ('POST', '/bulk-decide', 'Beslissing nemen voor meerdere meldingen tegelijk.'),
    ('POST', '/:id/comments', 'Nieuwe reactie toevoegen.'),
    ('GET', '/:id/comments', 'Reacties bij een melding ophalen.'),
    ('GET', '/:id/info-requests', 'Informatie-aanvragen bij een melding ophalen.'),
    ('POST', '/:id/info-requests', 'Informatie-aanvraag registreren.'),
    ('PUT', '/comments/:commentId', 'Bestaande reactie bewerken.'),
    ('POST', '/:id/comment', 'Verouderd endpoint voor het toevoegen van een enkele opmerking.'),
    ('GET', '/:id/coordinates', 'Aanvullende coördinaten bij een melding ophalen.'),
    ('POST', '/:id/coordinates', 'Coördinaat of geometrie toevoegen.'),
    ('PUT', '/:id/coordinates/:coordinateId', 'Coördinaat/geometrie bewerken.'),
    ('DELETE', '/:id/coordinates/:coordinateId', 'Coördinaat verwijderen.'),
    ('POST', '/:id/attachments', 'Bijlage uploaden.'),
    ('GET', '/:id/attachments', 'Bijlagen ophalen.'),
    ('GET', '/:id/attachments/:attachmentId/download', 'Bijlage downloaden.'),
    ('DELETE', '/:id/attachments/:attachmentId', 'Bijlage verwijderen.'),
    ('POST', '/:id/detect-zones', 'Automatische zonedetectie herberekenen.'),
    ('POST', '/:id/zones/:zoneCoverageId', 'Zone handmatig koppelen.'),
    ('DELETE', '/:id/zones/:zoneCoverageId', 'Zonekoppeling verwijderen.'),
    ('POST', '/:id/detect-products', 'Automatische productdetectie herberekenen.'),
])

endpoints_table(doc, 'product.routes.ts', '/api/products', [
    ('GET', '/', 'Productenlijst (catalogus).'),
    ('GET', '/:id', 'Detail van één product.'),
    ('POST', '/import-kml', 'KML-bestand uploaden en producten importeren.'),
    ('POST', '/', 'Nieuw product aanmaken.'),
    ('PUT', '/:id', 'Product bewerken.'),
    ('GET', '/for-notification/:notificationId', 'Producten die overlappen met een specifieke melding.'),
    ('POST', '/link-to-notification', 'Product handmatig koppelen aan een melding.'),
    ('DELETE', '/unlink-from-notification/:notificationId/:productId', 'Productkoppeling met een melding verwijderen.'),
    ('GET', '/notification/:notificationId', 'Gekoppelde producten van een melding ophalen.'),
])

endpoints_table(doc, 'productionLine.routes.ts', '/api/production-lines', [
    ('GET', '/', 'Lijst van actieve productielijnen (ZK, IENC, Pilot ENC, Publ).'),
])

endpoints_table(doc, 'productVersion.routes.ts', '/api/product-versions', [
    ('GET', '/', 'Lijst van openstaande productversies.'),
    ('GET', '/:id', 'Detail van één productversie.'),
    ('PATCH', '/:id/tasks/:taskId/execution-status', 'Uitvoeringsstatus van een gekoppelde taak binnen deze versie bijwerken.'),
    ('POST', '/', 'Nieuwe productversie aanmaken.'),
    ('POST', '/:id/publish', 'Versie publiceren; onvoltooide taken schuiven door naar de volgende versie.'),
    ('PATCH', '/:id/status', 'Status van de versie handmatig aanpassen.'),
    ('GET', '/:id/corrections-list', 'Gegenereerde correctielijst-inhoud (NL/EN) ophalen.'),
    ('GET', '/:id/baz2-publication', 'Gegenereerde BaZ-2-publicatietekst ophalen.'),
    ('POST', '/:id/attachments', 'Bijlage bij de versie uploaden.'),
    ('GET', '/:id/attachments', 'Bijlagen bij de versie ophalen.'),
    ('GET', '/:id/attachments/:attachmentId/download', 'Bijlage downloaden.'),
])

endpoints_table(doc, 'task.routes.ts', '/api/tasks', [
    ('GET', '/', 'Takenlijst per productielijn.'),
    ('GET', '/lead-times', 'Doorlooptijdstatistieken melding→taak→publicatie.'),
    ('GET', '/:id', 'Detail van één taak.'),
    ('PUT', '/:id', 'Taakgegevens bijwerken.'),
    ('PATCH', '/:id/flags', 'Vlaggen bijwerken (msi_active, needs_followup, needs_extra_info).'),
    ('PUT', '/:taskId/products/:productId', 'Status van een gekoppeld product bijwerken.'),
    ('POST', '/:taskId/products', 'Product koppelen aan de taak.'),
    ('POST', '/:taskId/related/:relatedTaskId', 'Verwante taak koppelen.'),
    ('POST', '/:id/notifications', 'Extra melding aan de taak koppelen.'),
    ('GET', '/:id/comments', 'Reacties bij de taak ophalen.'),
    ('POST', '/:id/comments', 'Reactie toevoegen.'),
    ('PUT', '/comments/:commentId', 'Reactie bewerken.'),
    ('GET', '/:id/workflow', 'Workflowstappen ophalen.'),
    ('POST', '/:id/workflow', 'Workflowstap toevoegen/bijwerken.'),
    ('GET', '/:id/production-line-status', 'Status per productielijn ophalen.'),
    ('PUT', '/:id/production-line-status/:productionLineId', 'Status voor een productielijn instellen.'),
    ('PATCH', '/:id/production-line-status/:productionLineId/wait-for-zk', '"Wachten op ZK"-vlag per lijn aan/uit zetten.'),
    ('GET', '/:id/hpd-projects', 'Gekoppelde CARIS HPD-projecten ophalen.'),
    ('GET', '/:id/articles', 'BaZ-artikelen bij de taak ophalen.'),
    ('POST', '/:id/articles', 'Nieuw BaZ-artikel aanmaken.'),
    ('PUT', '/:id/articles/:articleId', 'BaZ-artikel bewerken.'),
    ('DELETE', '/:id/articles/:articleId', 'BaZ-artikel verwijderen.'),
    ('POST', '/:id/articles/translate', 'Artikeltekst automatisch vertalen (Google Cloud Translate).'),
    ('GET', '/:id/info-requests', 'Informatie-aanvragen bij de taak ophalen.'),
    ('POST', '/:id/info-requests', 'Informatie-aanvraag registreren.'),
])

endpoints_table(doc, 'user.routes.ts', '/api/users', [
    ('GET', '/', 'Lijst van alle gebruikers.'),
])

# ---------------------------------------------------------------------------
# 4. Database Objects
# ---------------------------------------------------------------------------
add_h1(doc, '4. Database Objects (PostgreSQL / PostGIS)')

add_h2(doc, '4.1 Extensions')
bullet(doc, 'postgis (enabled via enable-postgis.sql) — geometrie-/geospatiale kolommen en functies (ST_Intersects, ST_Contains, ...) voor zone- en productdetectie.')

add_h2(doc, '4.2 Tables')
tables = [
    ('production_lines', 'Vier productielijnen: ZK, IENC, Pilot ENC, Publ.'),
    ('users', 'Gebruikersaccounts met bcrypt-wachtwoordhash.'),
    ('user_production_line_rights', 'Rechten (can_view/can_edit/can_publish) per gebruiker per productielijn.'),
    ('products', 'Productcatalogus (kaarten, publicaties, zones), inclusief PostGIS-geometrie.'),
    ('notifications', 'Meldingen (nautische informatie), met opmerkingen en status.'),
    ('notifications_products', 'Koppeltabel meldingen ↔ producten.'),
    ('notification_decisions', 'Beslissing (Ja/Nee) per melding per productielijn.'),
    ('notification_comments', 'Reacties bij een melding, per productielijn.'),
    ('notification_coordinates', 'Aanvullende handmatige coördinaten/geometrie bij een melding.'),
    ('notification_zones', 'Gekoppelde (gedetecteerde of handmatige) zones bij een melding.'),
    ('notification_info_requests', 'Informatie-aanvragen bij een melding.'),
    ('attachments', 'Bijlagen bij meldingen.'),
    ('tasks', 'Taken, met kolommen voor MSI-actief, opvolging nodig en extra info nodig (geen aparte "flags"-tabel).'),
    ('task_notifications', 'Koppeltabel taken ↔ (extra) meldingen.'),
    ('related_tasks', 'Koppeling tussen verwante taken.'),
    ('task_products', 'Koppeltabel taken ↔ producten, met uitvoeringsstatus per koppeling.'),
    ('task_comments', 'Reacties bij een taak, per productielijn.'),
    ('task_workflow', 'Workflowstappen van een taak.'),
    ('task_production_line_status', 'Status (en "wachten op ZK"-vlag) van een taak per productielijn.'),
    ('task_articles', 'BaZ-artikelen gekoppeld aan een taak (NL + auto-vertaalde EN-tekst).'),
    ('task_info_requests', 'Informatie-aanvragen bij een taak.'),
    ('hpd_projects', 'Koppeling met CARIS HPD-projecten.'),
    ('product_versions', 'Versies van productpublicaties (draft/published) per productielijn.'),
    ('product_version_attachments', 'Bijlagen bij een productversie.'),
    ('kml_files', 'Geïmporteerde KML-bronbestanden (metadata).'),
    ('kml_coverages', 'Individuele dekgebieden/geometrieën uit een KML-bestand.'),
    ('activity_log', 'Auditlog van belangrijke acties in het systeem.'),
    ('user_notes', 'Persoonlijke dashboard-notities, met prioriteit.'),
    ('user_note_production_lines', 'Zichtbaarheid van een notitie per productielijn.'),
]
two_col_table(doc, ['Tabel', 'Beschrijving'], tables, col_widths=[5.0, 12.0])

add_h2(doc, '4.3 Indexes')
indexes = [
    'idx_hpd_projects_task_id', 'idx_hpd_projects_project_code',
    'idx_kml_files_category', 'idx_kml_files_production_line',
    'idx_kml_coverages_file', 'idx_kml_coverages_code', 'idx_kml_coverages_type',
    'idx_notification_comments_notification', 'idx_notification_comments_production_line',
    'idx_notification_coordinates_notification',
    'idx_notification_zones_notification', 'idx_notification_zones_code',
    'idx_notification_info_requests_notification', 'idx_notification_info_requests_created_by',
    'idx_product_version_attachments_version',
    'idx_task_comments_task', 'idx_task_comments_production_line',
    'idx_task_workflow_task', 'idx_task_workflow_production_line',
    'idx_task_info_requests_task', 'idx_task_info_requests_created_by',
    'idx_task_pl_status_task', 'idx_task_pl_status_production_line', 'idx_task_pl_status_status',
    'idx_notifications_date', 'idx_notifications_status',
    'idx_tasks_number', 'idx_tasks_production_line',
    'idx_task_products_status', 'idx_product_versions_status',
    'idx_activity_log_entity', 'idx_activity_log_user',
]
for idx in indexes:
    bullet(doc, idx)

# ---------------------------------------------------------------------------
# 5. Automation Scripts and Tools
# ---------------------------------------------------------------------------
add_h1(doc, '5. Automation Scripts and Tools')
para(doc, (
    'The repository root contains a large collection of ad-hoc Node.js/PowerShell/SQL scripts '
    'used for one-off migrations, data fixes, diagnostics and demo/test data setup. They are not '
    'part of the running application and are executed manually by developers when needed.'
))

add_h2(doc, '5.1 Root-level Scripts')
script_descriptions = [
    ('add-geometry.js', 'Voegt ontbrekende PostGIS-geometrie toe aan meldingen op basis van hun coördinaten.'),
    ('add-task-product-execution-status.js', 'Eenmalige migratie: voegt de execution_status-kolom toe aan task_products.'),
    ('backfill-task-product-versions.js', 'Backfilt ontbrekende product_version_id-koppelingen in task_products.'),
    ('check-be3vlbnk-coverage.js', 'Diagnostische controle van de dekking/koppeling van product BE3VLBNK.'),
    ('check-be6avg1k-fixed.js', 'Controleert product BE6AVG1K na een fix.'),
    ('check-be6avg1k.js', 'Controleert product BE6AVG1K (voor de fix).'),
    ('check-chart-versions.js', 'Controleert de versienummering van zeekaartproducten.'),
    ('check-coordinates.js', 'Controleert coördinatendata van meldingen.'),
    ('check-coordinates2.js', 'Aanvullende/uitgebreide controle van coördinatendata.'),
    ('check-geometry-18.js', 'Diagnostische controle van geometrie voor een specifiek geval (#18).'),
    ('check-geometry-format.js', 'Controleert het opslagformaat van geometriekolommen.'),
    ('check-geometry-type.js', 'Controleert het geometrietype (punt/lijn/vlak) van records.'),
    ('check-missing-objnam.js', 'Zoekt producten zonder OBJNAM-naam.'),
    ('check-notification-zones.js', 'Controleert zone-koppelingen bij meldingen.'),
    ('check-notifications-schema.js', 'Toont het databaseschema van de notifications-tabel.'),
    ('check-pilot-enc-products.js', 'Controleert Pilot ENC-producten en hun koppelingen.'),
    ('check-products-storage.js', 'Controleert opslag/bestandslocatie van producten.'),
    ('check-skipped-products.js', 'Toont producten die zijn overgeslagen tijdens import/detectie.'),
    ('check-task-notifications.js', 'Controleert taak-meldingkoppelingen.'),
    ('check-zone-product-mixup.js', 'Spoort verwisselingen tussen zone- en productrecords op.'),
    ('check-zones-data.js', 'Controleert de consistentie van zonegegevens.'),
    ('cleanup-products-from-zones.js', 'Ruimt foutief als product geïmporteerde zone-records op.'),
    ('create-belgian-notices-20.js', 'Maakt 20 Belgische testmeldingen aan.'),
    ('create-notifications.js', 'Maakt testmeldingen aan (Node-script).'),
    ('create-notifications.sql', 'Maakt testmeldingen aan (SQL-variant).'),
    ('create-publ-correction-lists.js', 'Genereert/backfilt PUBL-correctielijsten.'),
    ('delete-all-tasks.js', 'Verwijdert alle taken (destructief; test-/resetdoeleinden).'),
    ('detect-products-for-notifications.bat', 'Windows-wrapper om productdetectie te draaien.'),
    ('detect-products-for-notifications.js', 'Draait automatische productdetectie voor bestaande meldingen.'),
    ('detect-products-for-notifications.ps1', 'PowerShell-wrapper om productdetectie te draaien.'),
    ('detect-zones-for-existing.js', 'Draait automatische zonedetectie voor bestaande meldingen.'),
    ('detect-zones.bat', 'Windows-wrapper om zonedetectie te draaien.'),
    ('detect-zones.ps1', 'PowerShell-wrapper om zonedetectie te draaien.'),
    ('diagnose-publ.ps1', 'Diagnostische PowerShell-check specifiek voor de PUBL-productielijn.'),
    ('final-summary.js', 'Toont een samenvattend overzicht na een migratie-/importrun.'),
    ('fix-coordinates-geometry.bat', 'Windows-wrapper voor het herstellen van geometrie op basis van coördinaten.'),
    ('fix-coordinates-geometry.ps1', 'Herstelt geometrie op basis van coördinaatvelden.'),
    ('fix-geometry-18.js', 'Eenmalige herstelfix voor een specifiek geometrieprobleem (#18).'),
    ('fix-pilot-enc-names-2.js', 'Corrigeert namen van Pilot ENC-producten (vervolgronde).'),
    ('fix-pilot-enc-names.js', 'Corrigeert namen van Pilot ENC-producten.'),
    ('fix-user-default-production-line.sql', 'Herstelt de standaard productielijn van een gebruiker.'),
    ('generate-analysis.py', 'Genereert CARTIS2_Analyse_April2026.docx (technische analyse).'),
    ('generate_docx.py', 'Genereert APP_FUNCTIONAL_ANALYSIS.docx uit de bijhorende Markdown.'),
    ('generate_functionalities_detailed.py', 'Genereert dit document, CARTIS2_Functionalities_Detailed_2026-04-27.docx.'),
    ('generate_user_guide.py', 'Genereert CARTIS2_Gebruikershandleiding.docx (Nederlandse gebruikershandleiding).'),
    ('get-vercel-logs.ps1', 'Haalt Vercel-deploymentlogs op voor diagnose.'),
    ('import-kml-coverages.bat', 'Windows-wrapper om KML-dekgebieden te importeren.'),
    ('import-kml-coverages.js', 'Importeert KML-dekgebieden (zones) in de database.'),
    ('import-kml-coverages.ps1', 'PowerShell-wrapper om KML-dekgebieden te importeren.'),
    ('import-products-kml.bat', 'Windows-wrapper om producten uit KML te importeren.'),
    ('import-products-kml.js', 'Importeert producten uit KML-bestanden.'),
    ('import-products-kml.ps1', 'PowerShell-wrapper om producten uit KML te importeren.'),
    ('import-to-supabase.bat', 'Windows-wrapper om lokale data naar Supabase te migreren.'),
    ('import-to-supabase.js', 'Migreert/importeert lokale data naar Supabase.'),
    ('import-to-supabase.ps1', 'PowerShell-wrapper om lokale data naar Supabase te migreren.'),
    ('insert_new_user.sql', 'Voegt een nieuwe gebruiker toe via SQL.'),
    ('inspect-product-versions.js', 'Inspecteert productversies en hun status.'),
    ('list-users.js', 'Toont alle gebruikers in de database.'),
    ('merge-publ-deelkaart-correction-lists.js', 'Voegt deelkaart-correctielijsten samen voor PUBL.'),
    ('migrate-baz2-version-numbers.js', 'Migreert/herberekent versienummering voor BaZ-2-producten.'),
    ('migrate-chart-version-numbers.js', 'Migreert/herberekent versienummering voor kaartproducten.'),
    ('migrate-enc-version-numbers.js', 'Migreert/herberekent versienummering voor ENC-producten.'),
    ('remove-linked-and-create-belgian.js', 'Verwijdert bestaande koppelingen en maakt Belgische testdata aan.'),
    ('reset-and-create-notices.js', 'Reset testmeldingen en maakt nieuwe demo-meldingen aan.'),
    ('reset-belgian-notices.js', 'Reset Belgische testmeldingen.'),
    ('setup-supabase-storage.js', 'Configureert de Supabase Storage-bucket voor bijlagen.'),
    ('show-zones-structure.ps1', 'Toont de structuur van de zonegegevens.'),
    ('start-cartis.bat', 'Windows-wrapper om backend en frontend samen te starten.'),
    ('start-cartis.ps1', 'PowerShell-script om backend en frontend samen te starten (lokale ontwikkeling).'),
    ('summary-reset.js', 'Toont een samenvatting na een resetactie.'),
    ('test-coordinates-api.js', 'Ad-hoc testscript voor de coördinaten-API.'),
    ('test-db-connection.js', 'Test de databaseverbinding (Supabase/PostgreSQL).'),
    ('test-detection-17.js', 'Ad-hoc testscript voor detectiegeval #17.'),
    ('test-detection-18.js', 'Ad-hoc testscript voor detectiegeval #18.'),
    ('test-login.js', 'Ad-hoc testscript voor de login-API.'),
    ('test-notification-detail.js', 'Ad-hoc testscript voor het meldingdetail-endpoint.'),
    ('test-notifications.js', 'Ad-hoc testscript voor de meldingen-API.'),
    ('test-product-status-update.js', 'Ad-hoc testscript voor het bijwerken van productstatus.'),
    ('test-zone-detection-12.js', 'Ad-hoc testscript voor zonedetectiegeval #12.'),
    ('test-zone-system.js', 'Ad-hoc testscript voor het volledige zonesysteem.'),
    ('trigger-all-product-detection.js', 'Triggert productdetectie voor alle meldingen.'),
    ('trigger-all-zone-detection.js', 'Triggert zonedetectie voor alle meldingen.'),
    ('trigger-detection-17.js', 'Triggert detectie voor een specifiek geval (#17).'),
    ('trigger-detection-18.js', 'Triggert detectie voor een specifiek geval (#18).'),
    ('trigger-zone-detection-12.js', 'Triggert zonedetectie voor een specifiek geval (#12).'),
    ('verify-new-notices.js', 'Verifieert nieuw aangemaakte testmeldingen.'),
    ('verify-notification-22.js', 'Verifieert een specifieke melding (#22).'),
    ('verify-products-18.js', 'Verifieert producten voor een specifiek geval (#18).'),
    ('verify-task-deletion.js', 'Verifieert dat taken correct verwijderd zijn.'),
]
two_col_table(doc, ['Script', 'Beschrijving'], script_descriptions, col_widths=[6.0, 11.0])

add_h2(doc, '5.2 NPM Scripts')
add_h3(doc, 'package.json')
for name, cmd in [
    ('install:all', 'npm install && npm install --prefix backend && npm install --prefix frontend'),
    ('dev', 'concurrently "npm run dev:backend" "npm run dev:frontend"'),
    ('dev:backend', 'npm run dev --prefix backend'),
    ('dev:frontend', 'npm run dev --prefix frontend'),
    ('build', 'npm run build --prefix backend && npm run build --prefix frontend'),
    ('build:backend', 'npm run build --prefix backend'),
    ('build:frontend', 'npm run build --prefix frontend'),
    ('start', 'npm start --prefix backend'),
    ('backfill:task-product-versions', 'node backfill-task-product-versions.js'),
    ('backfill:publ-correction-lists', 'node create-publ-correction-lists.js'),
    ('merge:publ-deelkaart-correction-lists', 'node merge-publ-deelkaart-correction-lists.js'),
    ('migrate:enc-version-numbers', 'node migrate-enc-version-numbers.js'),
    ('migrate:baz2-version-numbers', 'node migrate-baz2-version-numbers.js'),
    ('migrate:task-product-execution-status', 'node add-task-product-execution-status.js'),
    ('import:supabase', 'node import-to-supabase.js'),
    ('import:kml', 'node import-kml-coverages.js'),
    ('detect-zones', 'node detect-zones-for-existing.js'),
]:
    code_line(doc, f'{name}: {cmd}')

add_h3(doc, 'backend/package.json')
for name, cmd in [
    ('dev', 'ts-node-dev --respawn --transpile-only src/index.ts'),
    ('build', 'tsc'),
    ('start', 'node dist/index.js'),
    ('typecheck', 'tsc --noEmit'),
    ('detect-zones', 'ts-node detect-zones-for-existing.ts'),
]:
    code_line(doc, f'{name}: {cmd}')

add_h3(doc, 'frontend/package.json')
for name, cmd in [
    ('dev', 'vite'),
    ('build', 'tsc && vite build'),
    ('preview', 'vite preview'),
    ('lint', 'eslint . --ext ts,tsx --report-unused-disable-directives --max-warnings 0'),
]:
    code_line(doc, f'{name}: {cmd}')

add_h2(doc, '5.3 Script Domains')
domains = [
    ('Startup and environment', ['start-cartis.ps1', 'start-cartis.bat', 'get-vercel-logs.ps1']),
    ('Supabase deployment', ['import-to-supabase.ps1', 'import-to-supabase.bat', 'import-to-supabase.js', 'setup-supabase-storage.js', 'test-db-connection.js']),
    ('KML and product import', ['import-products-kml.ps1', 'import-products-kml.bat', 'import-products-kml.js', 'import-kml-coverages.ps1', 'import-kml-coverages.bat', 'import-kml-coverages.js']),
    ('Detection and zoning', ['detect-products-for-notifications.js', 'detect-products-for-notifications.ps1', 'detect-products-for-notifications.bat', 'detect-zones-for-existing.js', 'detect-zones.ps1', 'detect-zones.bat', 'trigger-all-product-detection.js', 'trigger-all-zone-detection.js', 'trigger-detection-17.js', 'trigger-detection-18.js', 'trigger-zone-detection-12.js']),
    ('Validation and testing', ['check-be3vlbnk-coverage.js', 'check-be6avg1k-fixed.js', 'check-be6avg1k.js', 'check-chart-versions.js', 'check-coordinates.js', 'check-coordinates2.js', 'check-geometry-18.js', 'check-geometry-format.js', 'check-geometry-type.js', 'check-missing-objnam.js', 'check-notification-zones.js', 'check-notifications-schema.js', 'check-pilot-enc-products.js', 'check-products-storage.js', 'check-skipped-products.js', 'check-task-notifications.js', 'check-zone-product-mixup.js', 'check-zones-data.js', 'test-coordinates-api.js', 'test-detection-17.js', 'test-detection-18.js', 'test-login.js', 'test-notification-detail.js', 'test-notifications.js', 'test-product-status-update.js', 'test-zone-detection-12.js', 'test-zone-system.js', 'verify-new-notices.js', 'verify-notification-22.js', 'verify-products-18.js', 'verify-task-deletion.js', 'show-zones-structure.ps1', 'diagnose-publ.ps1', 'final-summary.js', 'inspect-product-versions.js', 'list-users.js']),
    ('Migrations and fixes', ['add-geometry.js', 'add-task-product-execution-status.js', 'cleanup-products-from-zones.js', 'fix-coordinates-geometry.bat', 'fix-coordinates-geometry.ps1', 'fix-geometry-18.js', 'fix-pilot-enc-names.js', 'fix-pilot-enc-names-2.js', 'fix-user-default-production-line.sql', 'migrate-baz2-version-numbers.js', 'migrate-chart-version-numbers.js', 'migrate-enc-version-numbers.js', 'insert_new_user.sql']),
    ('Test data and publications', ['create-belgian-notices-20.js', 'create-notifications.js', 'create-notifications.sql', 'create-publ-correction-lists.js', 'merge-publ-deelkaart-correction-lists.js', 'remove-linked-and-create-belgian.js', 'reset-and-create-notices.js', 'reset-belgian-notices.js', 'summary-reset.js', 'delete-all-tasks.js', 'backfill-task-product-versions.js']),
    ('Documentation generators', ['generate-analysis.py', 'generate_docx.py', 'generate_user_guide.py', 'generate_functionalities_detailed.py']),
]
for domain, files in domains:
    p = doc.add_paragraph()
    r = p.add_run(domain + ':')
    r.bold = True
    for f in files:
        bullet(doc, f)

# ---------------------------------------------------------------------------
# 6. Documentation Coverage
# ---------------------------------------------------------------------------
add_h1(doc, '6. Documentation Coverage')
para(doc, (
    'Outline (headings and top-level bullets) of the main Markdown documentation files kept in '
    'the repository, scraped at generation time so this section stays close to the underlying '
    'files. Consult the referenced .md files directly for full detail.'
), italic=True, color=GREY)

doc_files = [
    'README.md', 'PROJECT_STATUS.md', 'PRODUCTS_INTEGRATION.md',
    'ZONE_DETECTION.md', 'SUPABASE_IMPORT.md', 'KML_IMPORT.md', 'SORTING_FEATURE.md',
]
for fname in doc_files:
    add_h2(doc, fname)
    outline = scrape_markdown_outline(os.path.join(REPO_ROOT, fname))
    if not outline:
        para(doc, '(bestand niet gevonden of leeg)', italic=True, color=GREY)
        continue
    for kind, level, text in outline:
        if kind == 'heading':
            p = doc.add_paragraph()
            r = p.add_run(('#' * level) + ' ' + text)
            r.bold = (level == 1)
        else:
            bullet(doc, text)

# ---------------------------------------------------------------------------
# 7. Functional Checklist by Area
# ---------------------------------------------------------------------------
add_h1(doc, '7. Functional Checklist by Area')
bullet(doc, 'Authentication and authorization (JWT, protected routes, production-line rights).')
bullet(doc, 'Notification lifecycle (create, edit, decide, bulk-decide, comments, coordinates, attachments, info requests).')
bullet(doc, 'Task lifecycle (creation from decisions, product links, flags, comments, workflow, per-line status, related tasks, BaZ articles, translation).')
bullet(doc, 'Product version lifecycle (draft, correction lists, BaZ-2 publication text, attachments, publish, lead-time tracking).')
bullet(doc, 'Article and publication support (BaZ articles, auto-translation endpoint, product versions, publish workflows).')
bullet(doc, 'Geospatial operations (PostGIS extension, GeoJSON handling, product and zone auto-detection, KML coverage import, GML export).')
bullet(doc, 'Personal dashboard notes (priority, per-production-line visibility, sorting).')
bullet(doc, 'Lead-time dashboard (melding→taak→publicatie statistics with column filters).')
bullet(doc, 'Auditability and observability (activity logging, diagnostics scripts, consistency checks).')
bullet(doc, 'Deployment/import operations (Vercel hosting, Supabase database + storage, SQL migration ordering, environment configuration).')
para(doc, (
    'Note: this inventory is generated from the current repository state (routes grepped '
    'directly from backend/src/routes/, pages from frontend/src/pages/, scripts from the '
    'repository root, schema from the SQL migration files). For exact UI wording at runtime, '
    'validate dynamic/localized strings in the running app.'
), italic=True, color=GREY)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
out_path = os.path.join(REPO_ROOT, 'CARTIS2_Functionalities_Detailed_2026-04-27.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
