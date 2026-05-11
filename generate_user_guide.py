"""Generate CARTIS 2.0 User Guide as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x5C)
MID_BLUE   = RGBColor(0x2E, 0x6D, 0xA4)
STEEL_BLUE = RGBColor(0x4A, 0x86, 0xC8)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x80, 0x80, 0x80)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_row(row, hex_color="D6E4F0"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'),  '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2E6DA4')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(4)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size      = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size      = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = STEEL_BLUE
        run.font.size      = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  ")
        r.bold = True
        r.font.color.rgb = DARK_BLUE
    p.add_run(text)
    return p

def note_box(doc, text, label="Opmerking"):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    shade_row(tbl.rows[0], "EAF3FB")
    p  = cell.paragraphs[0]
    rn = p.add_run(label + ":  ")
    rn.bold = True
    rn.font.color.rgb = MID_BLUE
    p.add_run(text)
    doc.add_paragraph()

def simple_table(doc, headers, rows, header_color="2E6DA4"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    shade_row(tbl.rows[0], header_color)
    for cell, txt in zip(tbl.rows[0].cells, headers):
        rn = cell.paragraphs[0].add_run(txt)
        rn.bold = True
        rn.font.color.rgb = WHITE
    for row_data in rows:
        row = tbl.add_row()
        for cell, val in zip(row.cells, row_data):
            cell.text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("CARTIS 2.0")
r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Gebruikershandleiding")
r2.font.size = Pt(22); r2.font.color.rgb = MID_BLUE

doc.add_paragraph()

org = doc.add_paragraph()
org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = org.add_run("Vlaamse Hydrografie")
r3.font.size = Pt(13); r3.font.italic = True; r3.font.color.rgb = STEEL_BLUE

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = ver.add_run("Versie 1.0  |  Mei 2026")
r4.font.size = Pt(11); r4.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "Inhoudsopgave")
toc = [
    ("1",    "Inleiding"),
    ("2",    "Aanmelden en gebruikersrechten"),
    ("3",    "Algemene navigatie"),
    ("4",    "Dashboard"),
    ("5",    "Meldingen"),
    ("5.1",  "Meldingen overzicht"),
    ("5.2",  "Nieuwe melding aanmaken"),
    ("5.3",  "Detailpagina van een melding"),
    ("6",    "Taken"),
    ("6.1",  "Takenoverzicht"),
    ("6.2",  "Detailpagina van een taak"),
    ("7",    "Productversies"),
    ("7.1",  "Openstaande productversies"),
    ("7.2",  "Gepubliceerde versies"),
    ("8",    "Producten"),
    ("9",    "Doorlooptijden"),
    ("10",   "Veelgestelde vragen en tips"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2 if "." not in num else 0.5)
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(11)
    if "." not in num:
        r.bold = True; r.font.color.rgb = DARK_BLUE
    else:
        r.font.color.rgb = MID_BLUE
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INLEIDING
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "1. Inleiding")
body(doc,
     "CARTIS 2.0 is de webapplicatie van Vlaamse Hydrografie voor de administratieve opvolging "
     "van nautische meldingen en de bijbehorende producten. Het systeem ondersteunt het volledige "
     "proces van melding-inname tot publicatie van nautische producten zoals zeekaarten, "
     "Inland ENC's, Pilot ENC's en publicaties (Berichten aan Zeevarenden, Lichtenlijst, "
     "Verbeterlijst, enz.).")

add_h2(doc, "1.1 Doelgroep")
body(doc,
     "Deze handleiding is bedoeld voor alle medewerkers van Vlaamse Hydrografie die dagelijks "
     "met CARTIS 2.0 werken: cartografen, redacteuren, teamleiders en beheerders.")

add_h2(doc, "1.2 Productielijnen")
body(doc,
     "CARTIS 2.0 organiseert het werk per productielijn. "
     "Elke gebruiker heeft per productielijn specifieke rechten.")
simple_table(doc,
    ["Productielijn", "Beschrijving"],
    [
        ("ZK",        "Zeekaartproductie – elektronische en papieren nautische kaarten."),
        ("IENC",      "Inland ENC – binnenvaartkaarten."),
        ("Pilot ENC", "Gedetailleerde bathymetrische loodskaarten."),
        ("Publ",      "Publicaties – BaZ, Lichtenlijst, Verbeterlijst, enz."),
    ],
    header_color="1A3A5C")

add_h2(doc, "1.3 Vier hoofdprocessen")
for num, proc, desc in [
    ("1", "Meldingen",      "Inkomende nautische informatie via API's, e-mail of handmatige invoer."),
    ("2", "Taken",          "Concrete werkitems gekoppeld aan meldingen, per productielijn."),
    ("3", "Productversies", "Versies van kaarten/publicaties waarin taken worden verwerkt."),
    ("4", "Publicaties",    "Het effectief publiceren van de nautische producten."),
]:
    bullet(doc, desc, bold_prefix=f"{num}. {proc}:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. AANMELDEN EN GEBRUIKERSRECHTEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "2. Aanmelden en gebruikersrechten")

add_h2(doc, "2.1 Aanmelden")
body(doc,
     "Navigeer naar de CARTIS 2.0 webapplicatie in uw browser. "
     "U krijgt automatisch de aanmeldpagina te zien als u nog niet ingelogd bent.")
for i, s in enumerate([
    "Voer uw e-mailadres in het veld 'E-mailadres' in.",
    "Voer uw wachtwoord in het veld 'Wachtwoord' in.",
    "Klik op de knop 'Inloggen'.",
    "Bij succesvolle aanmelding wordt u doorgestuurd naar het Dashboard.",
], 1):
    bullet(doc, s, bold_prefix=f"Stap {i}:")

note_box(doc,
    "Als uw account niet actief is of uw gegevens onjuist zijn, verschijnt een foutmelding "
    "bovenaan het aanmeldformulier. Neem contact op met de systeembeheerder als u niet kunt inloggen.")

add_h2(doc, "2.2 Uitloggen")
body(doc,
     "Klik rechts bovenaan op de knop 'Uitloggen' naast uw gebruikersnaam. "
     "Uw sessie wordt beëindigd en u wordt naar de aanmeldpagina doorverwezen.")

add_h2(doc, "2.3 Gebruikersrechten per productielijn")
body(doc, "Rechten zijn per productielijn ingesteld. U kunt de volgende rechten hebben:")
for right, desc in [
    ("Bekijken (can_view)",      "U kunt gegevens raadplegen maar niet wijzigen."),
    ("Bewerken (can_edit)",      "U kunt records aanmaken, bewerken en beslissingen nemen."),
    ("Publiceren (can_publish)", "U kunt productversies publiceren (omvat ook bewerken)."),
]:
    bullet(doc, desc, bold_prefix=right)
note_box(doc,
    "Sommige knoppen en acties zijn verborgen of uitgeschakeld als u onvoldoende rechten heeft. "
    "Neem contact op met uw beheerder voor rechteninstellingen.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ALGEMENE NAVIGATIE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "3. Algemene navigatie")

add_h2(doc, "3.1 Koptekstbalk (header)")
body(doc, "De koptekstbalk bovenaan het scherm bevat:")
for elem, desc in [
    ("App-logo / titel",     "Toont 'CARTIS 2.0'. Klikken brengt u terug naar het Dashboard."),
    ("Productielijn-selector","Kies de actieve productielijn via het dropdownmenu. De meeste pagina's zijn pas bruikbaar nadat u een productielijn heeft geselecteerd."),
    ("Gebruikersnaam",       "Toont de naam van de ingelogde gebruiker."),
    ("Uitloggen-knop",       "Beëindigt de sessie en stuurt u naar de aanmeldpagina."),
]:
    bullet(doc, desc, bold_prefix=elem)
note_box(doc,
    "Zolang er geen productielijn geselecteerd is, verschijnt een waarschuwingstekst en wordt "
    "het dropdownmenu gemarkeerd. Selecteer eerst een productielijn voor u verder werkt.")

add_h2(doc, "3.2 Zijbalk (sidebar)")
body(doc,
     "De zijbalk links bevat de navigatielinks naar alle hoofdsecties. "
     "U kunt de zijbalk in- en uitklappen via de schakelknop (☰ / ✕) linksboven.")
simple_table(doc,
    ["Sectie", "Beschrijving"],
    [
        ("Dashboard",            "Operationeel startscherm met KPI's en notities."),
        ("Meldingen",            "Overzicht en beheer van alle nautische meldingen."),
        ("Taken",                "Overzicht en beheer van alle taken per productielijn."),
        ("Productversies",       "Beheer van openstaande product-/publicatieversies."),
        ("Gepubliceerde versies","Archief van reeds gepubliceerde versies."),
        ("Producten",            "Productcatalogus (kaarten, publicaties) met kaartweergave."),
        ("Doorlooptijden",       "KPI-overzicht van doorlooptijden melding → taak → publicatie."),
    ])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "4. Dashboard")
body(doc,
     "Het Dashboard is het operationele startscherm na aanmelding. Het biedt een snel overzicht "
     "van de actuele werkdruk en dient als prikbord voor interne notities.")

add_h2(doc, "4.1 KPI-kaarten")
body(doc, "Bovenaan staan twee klikbare statistiekkaarten:")
for card, desc in [
    ("Openstaande Meldingen", "Aantal meldingen zonder definitieve beslissing. Klik om naar Meldingen te navigeren."),
    ("Actieve Taken",         "Aantal lopende taken voor de geselecteerde productielijn. Klik om naar Taken te navigeren."),
]:
    bullet(doc, desc, bold_prefix=card + ":")

add_h2(doc, "4.2 Notitiesbord")
body(doc,
     "Het notitiesbord toont interne werknotities die zichtbaar zijn per productielijn. "
     "U kunt notities aanmaken, bewerken, verbergen of verwijderen.")

add_h3(doc, "Sorteren")
body(doc, "Gebruik het sorteerdropdown rechtsboven:")
for opt in ["Prioriteit hoog → laag", "Prioriteit laag → hoog", "Nieuwste eerst", "Oudste eerst"]:
    bullet(doc, opt)

add_h3(doc, "Nieuwe notitie aanmaken")
for i, s in enumerate([
    "Klik op de knop 'Nieuwe nota'.",
    "Kies de prioriteit: laag, gemiddeld of hoog.",
    "Vink aan voor welke productielijnen de notitie zichtbaar moet zijn.",
    "Typ de inhoud in de teksteditor (opmaak is mogelijk).",
    "Klik op 'Nota opslaan'.",
], 1):
    bullet(doc, s, bold_prefix=f"Stap {i}:")

add_h3(doc, "Notities beheren")
for action, desc in [
    ("Bewerken",              "Opent de notitie in bewerkingsmodus (vereist bewerkerrechten)."),
    ("Houden (<lijn>)",       "Dwingt de notitie zichtbaar te houden voor de huidige productielijn."),
    ("Verwijderen voor lijn", "Verbergt de notitie enkel voor de huidige productielijn."),
    ("Verwijderen",           "Verwijdert de notitie definitief (enkel voor de aanmaker)."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. MELDINGEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "5. Meldingen")
body(doc,
     "De Meldingenmodule is het hart van CARTIS 2.0. Hier worden inkomende nautische meldingen "
     "ingevoerd, getrieerd en gekoppeld aan productielijnen en taken.")

# ── 5.1 ──────────────────────────────────────────────────────────────────────
add_h2(doc, "5.1 Meldingen overzicht")

add_h3(doc, "Filteropties")
for ctrl, desc in [
    ("Zoekbalk",         "Vrije tekstzoektocht over alle velden."),
    ("Scopefilter",      "Kies 'Alleen niet beslist' of 'Alle meldingen'."),
    ("Kolomfilters",     "Per kolom een filterinvoer (code, datum, titel, inhoud, zones, producten, taak, beslissing per lijn)."),
    ("Wis alle filters", "Verschijnt wanneer filters actief zijn; wist alle tegelijk."),
]:
    bullet(doc, desc, bold_prefix=ctrl + ":")

add_h3(doc, "Tabelkolommen")
simple_table(doc,
    ["Kolom", "Beschrijving"],
    [
        ("Selectievakje",      "Meerdere meldingen selecteren voor bulkacties."),
        ("Code",               "Unieke identificatiecode van de melding."),
        ("Datum",              "Datum waarop de melding ontvangen werd."),
        ("Titel",              "Korte beschrijving van de melding."),
        ("Inhoud (preview)",   "Samenvatting van de meldingsinhoud."),
        ("Zones",              "Gekoppelde geografische zones."),
        ("Producten",          "Automatisch of handmatig gekoppelde producten."),
        ("Taken",              "Gekoppelde taken."),
        ("Beslissing per lijn","Ja / Nee / Niet beslist per productielijn."),
        ("Acties",             "Details-knop, GML-export, beslissingsknop."),
    ])

add_h3(doc, "Rijacties")
for action, desc in [
    ("Details",       "Opent de detailpagina van de melding in een nieuw tabblad."),
    ("GML",           "Exporteert alle geometrieën als GML-bestand."),
    ("Ja / Nee",      "Beslissing voor de actieve productielijn (enkel zichtbaar als nog niet beslist)."),
    ("Uitvouwen (▶)", "Toont een uitgebreid detailpaneel direct in de tabel."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

add_h3(doc, "Bulkacties")
body(doc, "Selecteer meerdere rijen via de selectievakjes. Een bulkactiepaneel verschijnt met:")
for item in [
    "Een rijke teksteditor voor een gezamenlijke opmerking.",
    "'Ja voor alle X' – positieve beslissing voor alle geselecteerde meldingen.",
    "'Niet nodig voor alle X' – negatieve beslissing.",
    "Bij meerdere meldingen + 'Ja': keuze voor aparte taken per melding of één gemeenschappelijke taak.",
    "'Deselecteer alles' – maakt de selectie ongedaan.",
]:
    bullet(doc, item)

# ── 5.2 ──────────────────────────────────────────────────────────────────────
add_h2(doc, "5.2 Nieuwe melding aanmaken")
body(doc, "Klik op '+ Nieuwe Melding' rechtsboven om het aanmaakformulier te openen.")

add_h3(doc, "Basisvelden")
for fname, fdesc in [
    ("Titel *",         "Verplicht. Korte beschrijving van de melding."),
    ("Code",            "Optioneel referentienummer."),
    ("Bron",            "Selecteer de herkomstbron (bijv. Wrakkendatabank, RWS, e-mail, …)."),
    ("Bron Detail",     "Bijkomende brondetails (bijv. naam afzender of URL)."),
    ("Meldingsdatum *", "Verplicht. De datum waarop de melding ontvangen werd."),
    ("Inhoud",          "Gedetailleerde omschrijving van de melding."),
    ("Opmerkingen",     "Interne opmerkingen."),
]:
    bullet(doc, fdesc, bold_prefix=fname + ":")

add_h3(doc, "E-mail importeren")
body(doc,
     "Sleep een .eml- of .msg-bestand naar het importgebied of klik om een bestand te selecteren. "
     "CARTIS vult de velden automatisch in:")
for item in ["Afzender → Brondetail", "Onderwerp → Titel", "Berichttekst → Inhoud", "Bijlagen → bijlagelijst"]:
    bullet(doc, item)

add_h3(doc, "Bijlagen toevoegen")
body(doc,
     "Sleep bestanden naar het bijlagegebied of klik om bestanden te kiezen. "
     "Elke bijlage verschijnt in een lijst met een verwijderknop.")

add_h3(doc, "Locatie en geometrie")
body(doc, "Leg de geografische locatie op twee manieren vast:")
bullet(doc,
       "Kies Point, Lijn of Vlak via de tekenikonen en teken direct op de kaart.",
       bold_prefix="Tekenen op de kaart:")
bullet(doc,
       "Voer breedte- en lengtegraad in. Kies het coördinatenstelsel: "
       "DD, DDM, DMS, Lambert 72, Lambert 2008, ETRS89 UTM31N of WGS84 UTM31N.",
       bold_prefix="Coördinaten intypen:")
body(doc,
     "Voeg een optionele naam en beschrijving toe per geometrie. "
     "Klik 'Verwijder alles' om alle geometrieën te wissen.")

add_h3(doc, "WMS-lagen")
body(doc,
     "Klap het WMS-lagenpaneel open om maritieme referentiekaartlagen over de kaart te leggen "
     "als visuele hulp bij het intekenen.")

add_h3(doc, "Melding opslaan")
body(doc,
     "Klik op 'Aanmaken'. Verplichte velden zijn Titel en Meldingsdatum. "
     "Na aanmaak verschijnt de melding in het overzicht.")

# ── 5.3 ──────────────────────────────────────────────────────────────────────
add_h2(doc, "5.3 Detailpagina van een melding")
body(doc, "Klik op 'Details' in de actiekolom om de volledige detailpagina te openen.")

add_h3(doc, "Basisinformatie")
body(doc,
     "Bovenaan staan de basisgegevens: code, titel, bron, datum en inhoud. "
     "Klik op 'Bewerken' om de opmerkingen te bewerken en op 'Opslaan' om te bewaren.")

add_h3(doc, "Reacties (comments)")
body(doc,
     "Voeg productielijnspecifieke opmerkingen toe via de rijke teksteditor en klik op 'Verzenden'. "
     "Reacties worden per productielijn gegroepeerd.")

add_h3(doc, "Coördinaten en geometrie")
body(doc,
     "Voeg coördinaten toe, bewerk of verwijder ze. De kaart toont alle gekoppelde geometrieën. "
     "Gebruik de tekenikonen of het coördinatenformulier. "
     "De kaart ondersteunt volledig scherm via 'Vergroot kaart'.")

add_h3(doc, "Zones")
body(doc, "Zones zijn geografische gebieden die met de melding overlappen.")
for action, desc in [
    ("Automatische detectie",    "Klik op 'Zones detecteren' om zones automatisch te berekenen op basis van de geometrieën."),
    ("Handmatig toevoegen",      "Kies een zone uit het dropdown en klik op 'Toevoegen'."),
    ("Verwijderen",              "Klik op het 'x'-pictogram naast een zone om de koppeling te verwijderen."),
    ("Gebieden tonen/verbergen", "Schakel de zichtbaarheid van zone-gebieden in/uit op de kaart."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

add_h3(doc, "Producten")
body(doc, "Het productenpaneel toont welke nautische producten overlappen met de melding.")
for action, desc in [
    ("Detectie",        "Klik op 'Producten detecteren' voor automatische detectie."),
    ("Handmatig koppelen", "Zoek een product en klik op 'Koppelen'."),
    ("Ontkoppelen",     "Klik op 'Ontkoppelen' bij een gekoppeld product."),
    ("Kaartweergave",   "Schakel producttypen in/uit op de kaart."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

add_h3(doc, "Bijlagen")
body(doc,
     "Upload extra bestanden via het bijlagegebied. "
     "Bestaande bijlagen kunnen worden gedownload of verwijderd.")

add_h3(doc, "Informatieverzoek")
body(doc, "Maak een e-mailontwerp aan om aanvullende informatie op te vragen:")
for item in [
    "Vul de ontvanger, het onderwerp en de berichttekst in.",
    "Klik op 'Opslaan en openen' om het ontwerp te bewaren en uw e-mailclient te openen.",
]:
    bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. TAKEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "6. Taken")
body(doc,
     "Taken zijn concrete werkitems die voortvloeien uit meldingen. Elke taak is gekoppeld aan "
     "een of meerdere meldingen en heeft per productielijn een eigen status.")

add_h2(doc, "6.1 Takenoverzicht")

add_h3(doc, "Filteropties")
for ctrl, desc in [
    ("Zoekbalk",                    "Vrije tekstzoektocht."),
    ("Statusfilter (huidige lijn)", "Filter op status voor de actieve productielijn."),
    ("Statusfilter (andere lijnen)","Filter op status van een andere productielijn."),
    ("Kolomfilters wissen",         "Wist alle actieve kolomfilters."),
]:
    bullet(doc, desc, bold_prefix=ctrl + ":")

add_h3(doc, "Tabelkolommen")
simple_table(doc,
    ["Kolom", "Beschrijving"],
    [
        ("Taaknummer",            "Uniek identificatienummer."),
        ("Titel",                 "Korte taakomschrijving."),
        ("BaZ-nummer(s)",         "Gekoppelde Berichten aan Zeevarenden-nummers."),
        ("MSI",                   "Indicator voor Maritime Safety Information."),
        ("Opvolging",             "Vlag voor meldingen die opvolging vereisen."),
        ("Extra info",            "Vlag voor meldingen waarvoor extra informatie gevraagd is."),
        ("Producten",             "Gekoppelde producten (klikbaar naar productversies)."),
        ("Status huidige lijn",   "Status voor de actieve productielijn."),
        ("Wachten op ZK",         "Checkbox: taak wacht op zeekaartproductie (lijnspecifiek)."),
        ("Status andere lijnen",  "Overzicht van statussen voor overige productielijnen."),
        ("Acties",                "Detailknop opent de volledige taakpagina in een nieuw tabblad."),
    ])

add_h3(doc, "Inline acties")
for action, desc in [
    ("Opvolging-checkbox",    "Schakelt de opvolgingsvlag aan/uit."),
    ("Extra info-checkbox",   "Schakelt de extra-infovlag aan/uit."),
    ("Wachten op ZK",         "Schakelt de wacht-op-ZK-vlag aan/uit (lijnspecifiek)."),
    ("Rij uitvouwen",         "Toont een detailpaneel met omschrijving, BaZ-chips en het e-mailformulier voor meer info."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

add_h2(doc, "6.2 Detailpagina van een taak")
body(doc, "Klik op 'Details' om de volledige taakpagina te openen.")

add_h3(doc, "Taakinformatie en status")
body(doc,
     "Bovenaan staan de taakmeta-gegevens: nummer, titel, aanmaakdatum, gekoppelde meldingen "
     "en een omschrijving. Per productielijn kunt u de status instellen via het statusdropdown.")

add_h3(doc, "Reacties per productielijn")
body(doc,
     "Voeg opmerkingen toe per productielijn via de rijke teksteditor. "
     "Bestaande reacties kunnen worden bewerkt.")

add_h3(doc, "Workfloweditie")
body(doc,
     "Voor elke productielijn kan een werkstroom worden gedefinieerd met stappen en verantwoordelijken.")

add_h3(doc, "Producten beheren")
for action, desc in [
    ("Product koppelen",   "Selecteer productielijn en product, klik 'Koppelen'."),
    ("Uitvoeringsstatus",  "Stel per product de uitvoeringsstatus in (bijv. 'In uitvoering', 'Klaar')."),
    ("Kaartweergave",      "Bekijk de geometrieën van gekoppelde producten op de kaart."),
]:
    bullet(doc, desc, bold_prefix=action + ":")

add_h3(doc, "Meldingen koppelen")
body(doc, "Zoek en koppel bijkomende meldingen aan de taak via het meldingenzoekformulier.")

add_h3(doc, "BaZ-artikelen (publicatieworkflow)")
body(doc, "Voor de productielijn 'Publ' kunnen BaZ-artikelen worden beheerd:")
for item in [
    "Klik op 'Nieuw artikel' om een artikel aan te maken.",
    "Vul het boeknummer, de Nederlandse titel en inhoud in.",
    "Gebruik 'Vertalen' om automatisch een Engelse versie te genereren.",
    "Sla het artikel op of verwijder het via de acties in de artikellijst.",
]:
    bullet(doc, item)

add_h3(doc, "Informatieverzoek")
body(doc,
     "Stel een e-mailontwerp op om aanvullende informatie op te vragen "
     "(zelfde werking als bij meldingen).")

add_h3(doc, "Kaart en splitscherm")
body(doc,
     "De kaart rechts toont producten, geometrieën en WMS-lagen. Gebruik 'Vergroot kaart' voor "
     "volledig scherm of versleep de splitbalk om de kaartbreedte aan te passen.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. PRODUCTVERSIES
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "7. Productversies")
body(doc,
     "In deze module beheert u de versies van nautische producten (kaarten, publicaties) "
     "die in productie zijn of reeds gepubliceerd zijn.")

add_h2(doc, "7.1 Openstaande productversies")

add_h3(doc, "Nieuwe productversie aanmaken")
body(doc, "Klap het paneel 'Nieuwe Productversie' open en vul de volgende velden in:")
for fname, fdesc in [
    ("Product *",       "Selecteer het product waarvoor u een versie aanmaakt."),
    ("Editienummer *",  "Voor handmatige producten: het editienummer."),
    ("Updatenummer *",  "Voor handmatige producten: het updatenummer."),
    ("Versiedatum *",   "De datum van de versie."),
    ("Notities",        "Optionele interne notities."),
]:
    bullet(doc, fdesc, bold_prefix=fname + ":")
note_box(doc,
    "Voor automatisch genummerde producten (BaZ-2, Lichtenlijst, correctielijsten) "
    "wordt het versienummer automatisch berekend en vooraf ingevuld.")
body(doc, "Klik op 'Versie aanmaken' om de versie op te slaan.")

add_h3(doc, "Versies tabel")
body(doc,
     "De tabel toont alle openstaande versies met: productcode, versienummer, datum, status, "
     "aanmaker en notities. Klik op een rij om het detailpaneel te openen.")

add_h3(doc, "Detailpaneel van een versie")
for block, desc in [
    ("Correctielijst preview", "Voor correctielijstproducten: bekijk de NL of EN versie en print als PDF via 'Print A4 (PDF)'."),
    ("Bijlagen",               "Upload brondocumenten bij de versie. Bestaande bijlagen kunnen worden gedownload."),
    ("Gekoppelde taken",       "Overzicht van taken in deze versie met uitvoeringsstatus. Pas de uitvoeringsstatus inline aan."),
    ("Publicatiecontroles",    "Vink 'Nieuwe editie' aan indien van toepassing. Voor kaartproducten is een publicatiedatum verplicht. Klik 'Publiceren' om de versie te publiceren."),
]:
    bullet(doc, desc, bold_prefix=block + ":")

add_h2(doc, "7.2 Gepubliceerde versies")
body(doc, "De pagina 'Gepubliceerde versies' is een archief van alle reeds gepubliceerde productversies.")

add_h3(doc, "Filteren")
body(doc,
     "Gebruik de kolomfilterinvoervelden boven de tabel om te filteren op product, versienummer, "
     "datum, publicatiedatum, aanmaker of notities. Klik 'Kolomfilters wissen' om alle filters te wissen.")

add_h3(doc, "Detailpaneel")
body(doc, "Klik op een rij om het detailpaneel te openen:")
for item in [
    "Gekoppelde taken met hun status en uitvoeringsstatus.",
    "Voor correctielijstproducten: NL/EN preview en printknop.",
    "BaZ-artikelpreview (NL + EN naast elkaar) via de knop in de takenlijst.",
]:
    bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. PRODUCTEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "8. Producten")
body(doc,
     "De Productencatalogus bevat alle nautische producten (kaarten, publicaties) die in "
     "CARTIS 2.0 worden beheerd. U kunt producten bekijken, bewerken en aanmaken (mits rechten).")

add_h2(doc, "8.1 Tabelweergave")
body(doc, "Klik op de knop 'Tabel' om de productlijst als tabel te bekijken.")
simple_table(doc,
    ["Kolom", "Beschrijving"],
    [
        ("Code",        "Unieke productcode."),
        ("Naam",        "Productnaam."),
        ("Beschrijving","OBJNAM-gebaseerde omschrijving."),
        ("Type",        "Producttype (bijv. ENC, kaart, publicatie)."),
        ("Actief",      "Geeft aan of het product actief in gebruik is."),
        ("Geometrie",   "Geeft aan of er een geometrie beschikbaar is."),
        ("Acties",      "Bewerkingsknop (vereist bewerkerrechten)."),
    ])

add_h2(doc, "8.2 Kaartweergave")
body(doc, "Klik op de knop 'Kaart' voor een geografische weergave van alle producten met geometrie.")
for ctrl, desc in [
    ("Categoriefilters",        "Filter op productcategorie (ENC Ux, IENC, Pilot ENC, Zeekaarten). Gebruik 'Alles selecteren'/'Alles deselecteren'."),
    ("Vergroot/verklein kaart", "Schakel een grotere kaartweergave in."),
    ("Productklikpopup",        "Klik op een product op de kaart voor metadata (code, naam, type)."),
    ("Legende",                 "De legende rechts groepeert producten per bestandscategorie."),
]:
    bullet(doc, desc, bold_prefix=ctrl + ":")

add_h2(doc, "8.3 Product aanmaken")
body(doc, "Klik op '+ Nieuw Product' (vereist bewerkerrechten). Vul de volgende velden in:")
for fname, fdesc in [
    ("Code *",       "Unieke productcode."),
    ("Naam *",       "Productnaam."),
    ("Type *",       "Selecteer het producttype."),
    ("Beschrijving", "Optionele omschrijving."),
]:
    bullet(doc, fdesc, bold_prefix=fname + ":")
body(doc, "Klik op 'Opslaan' om het product aan te maken.")

add_h2(doc, "8.4 Product bewerken")
body(doc,
     "Klik op 'Bewerken' in de actiekolom. Pas de naam, het type, de beschrijving of de "
     "actief-status aan en klik op 'Opslaan'.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. DOORLOOPTIJDEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "9. Doorlooptijden")
body(doc,
     "De Doorlooptijdenpagina geeft inzicht in de tijdsduur van het volledige traject: "
     "van melding-ontvangst tot publicatie van het bijbehorende product.")

add_h2(doc, "9.1 Samenvattingskaarten")
for card, desc in [
    ("Aantal trajecten",         "Totaal aantal afgeronde melding→taak→publicatie-trajecten."),
    ("Gem. melding → taak",      "Gemiddeld aantal dagen tussen melding-ontvangst en taakafsluiting."),
    ("Gem. taak → publicatie",   "Gemiddeld aantal dagen tussen taakafsluiting en publicatiedatum."),
    ("Gem. totale doorlooptijd", "Gemiddeld totaal aantal dagen van melding tot publicatie."),
]:
    bullet(doc, desc, bold_prefix=card + ":")

add_h2(doc, "9.2 Detailtabel")
body(doc, "De tabel toont alle afzonderlijke trajecten:")
simple_table(doc,
    ["Kolom", "Beschrijving"],
    [
        ("Melding",                 "Link naar de melding."),
        ("Ontvangstdatum",          "Datum en tijdstip van de melding."),
        ("Taak",                    "Link naar de taak."),
        ("Afsluitdatum taak",       "Datum en tijdstip van taakafsluiting."),
        ("Product",                 "Gekoppeld product."),
        ("Versie",                  "Productversie."),
        ("Publicatiedatum",         "Datum van publicatie."),
        ("Melding → taak (dagen)",  "Doorlooptijd in dagen."),
        ("Taak → publ. (dagen)",    "Doorlooptijd in dagen."),
        ("Totaal (dagen)",          "Totale doorlooptijd."),
    ])

add_h2(doc, "9.3 Filteren en zoeken")
body(doc,
     "Gebruik de zoekbalk bovenaan voor een vrije tekstzoektocht. "
     "Elke kolom heeft een filterinvoer. Klik 'Kolomfilters wissen' om alle kolomfilters te wissen.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. VEELGESTELDE VRAGEN EN TIPS
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "10. Veelgestelde vragen en tips")

faq = [
    ("Ik zie geen gegevens op de pagina.",
     "Controleer of u een productielijn heeft geselecteerd in het dropdownmenu bovenaan."),
    ("De knop 'Publiceren' is uitgeschakeld.",
     "U heeft 'publiceren'-rechten nodig voor de actieve productielijn. "
     "Neem contact op met uw beheerder. "
     "Controleer ook of een publicatiedatum is ingevuld (voor kaartproducten)."),
    ("Ik kan een melding niet bewerken.",
     "Controleer uw rechten (bewerken) voor de actieve productielijn. "
     "Zonder 'can_edit'-recht zijn bewerkingsknoppen verborgen."),
    ("Hoe importeer ik een e-mail als melding?",
     "Klik op '+ Nieuwe Melding', sleep uw .eml- of .msg-bestand naar het importgebied, "
     "en CARTIS vult de velden automatisch in."),
    ("Hoe voeg ik een geometrie toe aan een melding?",
     "Op de aanmaak- of detailpagina kunt u een geometrie tekenen op de kaart "
     "of handmatig coördinaten invoeren. "
     "Meerdere formaten worden ondersteund (DD, DDM, DMS, Lambert, …)."),
    ("Zones worden niet automatisch herkend.",
     "Klik op 'Zones detecteren' op de detailpagina van de melding. "
     "Er moet minstens één geometrie aanwezig zijn voor de detectie."),
    ("Hoe vertaal ik een BaZ-artikel naar het Engels?",
     "Open de taakdetailpagina, ga naar het BaZ-artikelenpaneel, "
     "open het artikel en klik op 'Vertalen'. De Engelse tekst wordt automatisch ingevuld."),
    ("Kan ik meerdere meldingen tegelijk afhandelen?",
     "Ja. Selecteer meerdere meldingen via de selectievakjes "
     "en gebruik de bulkacties (Ja voor alle / Niet nodig voor alle)."),
    ("Hoe druk ik een correctielijst af?",
     "Ga naar Productversies, selecteer de correctielijstversie en klik op 'Print A4 (PDF)' "
     "in het correctielijst-previewpaneel."),
    ("Wat is het verschil tussen 'Verwijderen voor lijn' en 'Verwijderen' bij notities?",
     "'Verwijderen voor lijn' verbergt de notitie enkel voor uw productielijn. "
     "'Verwijderen' wist de notitie definitief voor alle gebruikers (enkel voor de aanmaker)."),
    ("Hoe pas ik de uitvoeringsstatus van een product in een versie aan?",
     "Ga naar Productversies, selecteer de versie, scroll naar 'Gekoppelde taken' "
     "en wijzig de uitvoeringsstatus inline in de tabel."),
    ("Ik word automatisch uitgelogd.",
     "Uw sessietoken is verlopen (na een periode van inactiviteit). "
     "Log opnieuw in. Uw niet-opgeslagen gegevens gaan verloren."),
]

for q, a in faq:
    add_h3(doc, "V: " + q)
    body(doc, "A: " + a)

add_hr(doc)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("CARTIS 2.0 – Gebruikershandleiding  |  Vlaamse Hydrografie  |  Mei 2026")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"d:\Programming\Webapps\Cartis_new\CARTIS2_Gebruikershandleiding.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
