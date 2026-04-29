"""
Generate CARTIS 2.0 Analysis Document (DOCX)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    """Add a thin horizontal line (paragraph with bottom border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8C)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return h

def para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # light grey background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def two_col_table(doc, rows, header=None, col_widths=None):
    """Generic two-column table."""
    cols = 2
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    table.style = 'Table Grid'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    r_offset = 0
    if header:
        for i, h in enumerate(header):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            set_cell_bg(cell, '1A568C')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r_offset = 1
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + r_offset].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1:
                set_cell_bg(cell, 'EEF4FB')
    return table

# ─── Document ───────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── Title Page ─────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('CARTIS 2.0')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8C)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Analyse & Technische Documentatie')
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
doc.add_paragraph()

org_p = doc.add_paragraph()
org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = org_p.add_run('Vlaamse Hydrografie')
run3.bold = True
run3.font.size = Pt(13)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = date_p.add_run(f'Versie: April 2026')
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ─── 1. Inleiding ────────────────────────────────────────────────────────────

heading(doc, '1. Inleiding & Projectcontext', level=1)

para(doc, (
    'CARTIS (Cartografisch Informatiesysteem) is de applicatie van Vlaamse Hydrografie voor '
    'de administratieve opvolging van nautische meldingen en de daaraan gekoppelde producten. '
    'Het bestaande systeem is een module binnen KIS (Kust Informatie Systeem), opgezet in 2014–2015. '
    'Door de loskoppeling van KIS en de snelle digitale evolutie ontstond de nood aan een moderne, '
    'performante en sterk geautomatiseerde opvolger: CARTIS 2.0.'
))

heading(doc, '1.1 Doelstellingen', level=2)
bullet(doc, 'Meer automatisering, dataverwerking en data-uitwisseling (intern en extern).')
bullet(doc, 'REST-API\'s en semi-automatische datavisualisatie/-verwerking.')
bullet(doc, 'Automatische opmaak van nautische producten.')
bullet(doc, 'Ondersteuning voor diverse bestandsformaten (import/export).')
bullet(doc, 'Webapplicatie als front-end, met oog op migratie naar cloud.')
bullet(doc, 'Enkel cartografie-gerelateerde functionaliteiten (overige KIS-functies blijven buiten scope).')

heading(doc, '1.2 Geraamde schaal', level=2)
para(doc, (
    'Jaarlijks worden 2000–3000 nautische berichten verwerkt, afkomstig uit meerdere bronnen '
    '(API\'s, mailbox, handmatige invoer, pushberichten). Het systeem ondersteunt meerdere gelijktijdige '
    'gebruikers per productielijn.'
))

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 2. Architectuur ─────────────────────────────────────────────────────────

heading(doc, '2. Systeemarchitectuur', level=1)

para(doc, (
    'CARTIS 2.0 is opgezet als een monorepo met twee afzonderlijke sub-projecten: '
    'een Node.js/Express back-end en een React/TypeScript front-end, die communiceren via een REST API.'
))

heading(doc, '2.1 Technische stack', level=2)

two_col_table(doc,
    rows=[
        ('Back-end runtime',    'Node.js 20 + TypeScript 5'),
        ('Back-end framework',  'Express 4.18'),
        ('Database',            'PostgreSQL (met PostGIS extensie)'),
        ('ORM / Query',         'pg (node-postgres) — parameterized queries'),
        ('Authenticatie',       'JWT (jsonwebtoken) + bcryptjs'),
        ('Beveiliging',         'Helmet, CORS, express-validator'),
        ('File upload',         'Multer (max 10 MB; PDF, Office, afbeeldingen, XML, ZIP)'),
        ('Mail',                'Nodemailer (nautinfo mailbox – config aanwezig)'),
        ('Vertaling',           'Google Cloud Translate (@google-cloud/translate)'),
        ('Front-end framework', 'React 18 + TypeScript 5'),
        ('Build tool',          'Vite 5'),
        ('Routing',             'React Router DOM 6'),
        ('State management',    'Zustand (auth) + TanStack React Query 5 (server state)'),
        ('Kaarten',             'Leaflet 1.9 + react-leaflet + leaflet-draw'),
        ('Coördinaten',         'proj4 (7 coördinaatformaten: DD, DDM, DMS + 4 projected CRS)'),
        ('WYSIWYG editor',      'SunEditor (meldingen) / QuillEditor'),
        ('Export',              'JSZip + @mapbox/shp-write (Shapefile export)'),
        ('Logging',             'Morgan (HTTP logging)'),
    ],
    header=['Laag / Aspect', 'Technologie'],
    col_widths=[2.8, 4.0],
)

heading(doc, '2.2 Mappenstructuur', level=2)

code_block(doc, 'cartis-2.0/')
code_block(doc, '├── backend/                 # Node.js Express API')
code_block(doc, '│   ├── src/')
code_block(doc, '│   │   ├── config/         # DB- en omgevingsconfiguratie')
code_block(doc, '│   │   ├── middleware/     # Auth-middleware, foutafhandeling')
code_block(doc, '│   │   ├── routes/         # 8 route-modules (70+ endpoints)')
code_block(doc, '│   │   ├── services/       # Zone-detectie, HPD, productversies')
code_block(doc, '│   │   └── index.ts        # Express entry-point')
code_block(doc, '│   ├── database/           # schema.sql + 13 migratiescripts')
code_block(doc, '│   └── uploads/            # Geüploade bijlagen (lokale opslag)')
code_block(doc, '├── frontend/               # React + TypeScript')
code_block(doc, '│   └── src/')
code_block(doc, '│       ├── components/    # Layout, FileUpload, CoordinateInput, …')
code_block(doc, '│       ├── pages/         # 9 paginacomponenten')
code_block(doc, '│       ├── services/      # API-aanroepen')
code_block(doc, '│       ├── stores/        # Zustand auth-store')
code_block(doc, '│       ├── hooks/         # useTableSort, …')
code_block(doc, '│       └── utils/         # Hulpfuncties')
code_block(doc, '└── package.json            # Monorepo root (npm run dev start beide)')

heading(doc, '2.3 Communicatiemodel', level=2)

para(doc, (
    'De front-end (poort 5173 in dev) communiceert uitsluitend via HTTP(S) met de back-end REST API '
    '(poort 3000). Authenticatie verloopt via een Bearer JWT-token dat bij elke request meegestuurd wordt. '
    'De back-end beheert de verbinding met PostgreSQL via een connection pool.'
))

bullet(doc, 'Frontend → REST API (JSON over HTTP, Bearer token)')
bullet(doc, 'REST API → PostgreSQL (parameterized queries via pg pool)')
bullet(doc, 'REST API → Google Translate API (via @google-cloud/translate)')
bullet(doc, 'REST API → CARIS HPD (HPD Project Service, synchronisatie)')
bullet(doc, 'REST API → Filesystem (uploads/ map voor bijlagen)')

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 3. Database ─────────────────────────────────────────────────────────────

heading(doc, '3. Databaseschema', level=1)

para(doc, (
    'De PostgreSQL database (met PostGIS extensie voor ruimtelijke queries) bevat meer dan 20 tabellen. '
    'Hieronder volgt een overzicht van de kernentiteiten en hun onderlinge relaties.'
))

heading(doc, '3.1 Kernentiteiten', level=2)

two_col_table(doc,
    rows=[
        ('production_lines',            'Productielijnen: ZK, IENC, PILOT_ENC, PUBL'),
        ('users',                       'Gebruikers met e-mail, wachtwoord-hash, standaard productielijn'),
        ('user_production_line_rights', 'Rol-gebaseerde rechten (can_view / can_edit / can_publish) per gebruiker per productielijn'),
        ('notifications',               'Nautische meldingen: code, titel, inhoud, bron, datum, GeoJSON-geometrie, metadata (JSONB)'),
        ('notification_decisions',      'Beslissing (–/Ja/Nee) per melding per productielijn'),
        ('notification_coordinates',    'Aanvullende coördinaten bij een melding'),
        ('notification_comments',       'Commentaren bij meldingen'),
        ('notification_zones',          'Gedetecteerde geografische zones bij een melding (automatisch + manueel)'),
        ('notification_flags',          'Vlaggen: MSI actief, opvolging nodig, extra info nodig'),
        ('products',                    'Nautische producten (kaarten, publicaties) met GeoJSON-geometrie en productielijkoppeling'),
        ('notifications_products',      'M2M koppeling melding ↔ product (met is_relevant vlag)'),
        ('tasks',                       'Taken (taaknummer YYNNN, titel, BaZ-nr, CARIS HPD-koppeling)'),
        ('task_notifications',          'M2M koppeling taak ↔ melding'),
        ('task_products',               'Koppeling taak ↔ product met status-tracking'),
        ('task_comments',               'Commentaren bij taken'),
        ('task_workflow',               'Workflow-statussen per taak'),
        ('task_production_line_status', 'Productielijn-specifieke status per taak'),
        ('task_articles',               'BaZ/publicatie-artikelen bij taken (meertalig via Google Translate)'),
        ('product_versions',            'Versies van een product (status: In Progress / Ready / Published)'),
        ('related_tasks',               'M2M koppeling verwante taken'),
        ('attachments',                 'Bijlagen bij meldingen (bestandsnaam, pad, MIME-type, grootte)'),
        ('kml_coverages',               'Geïmporteerde KML-dekgebieden (zones)'),
        ('hpd_projects',                'CARIS HPD-projectkoppelingen'),
        ('activity_log',                'Audit trail van alle acties'),
    ],
    header=['Tabel', 'Beschrijving'],
    col_widths=[2.5, 4.3],
)

heading(doc, '3.2 Ruimtelijke functionaliteit (PostGIS)', level=2)
bullet(doc, 'PostGIS extensie voor geometrieverwerking en ruimtelijke queries.')
bullet(doc, 'geojson_intersects() functie voor overlap-detectie tussen producten en meldingen.')
bullet(doc, 'Ruimtelijke indexen op geometry-kolommen van products en notifications.')
bullet(doc, 'Ray-casting algoritme (TypeScript) voor point-in-polygon zone-detectie.')

heading(doc, '3.3 Migratiegeschiedenis', level=2)
para(doc, 'Het schema is uitgebreid via 13 incrementele SQL-migratiescripts:')

migration_rows = [
    ('add-geometry-to-coordinates.sql',         'Geometrie-kolom aan coördinaten-tabel'),
    ('add-hpd-projects.sql',                    'CARIS HPD-projectkoppeling'),
    ('add-kml-coverages.sql',                   'KML-dekgebieden (zones)'),
    ('add-notification-comments.sql',           'Commentaren bij meldingen'),
    ('add-notification-coordinates.sql',        'Aanvullende coördinaten'),
    ('add-notification-flags.sql',              'Vlaggen (MSI, opvolging, extra info)'),
    ('add-notification-zones.sql',              'Zone-koppelingen bij meldingen'),
    ('add-opmerkingen.sql',                     'Opmerkingenveld'),
    ('add-task-articles.sql',                   'Artikelen bij taken'),
    ('add-task-comments-and-workflow.sql',      'Taakcommentaren en workflow'),
    ('add-task-info-requests.sql',              'Informatie-aanvragen bij taken'),
    ('add-task-production-line-status.sql',     'Productielijn-status per taak'),
    ('add-article-titles.sql',                  'Titels voor artikelen'),
]
two_col_table(doc, migration_rows, header=['Migratiescript', 'Doel'], col_widths=[3.2, 3.6])

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 4. Back-end API ─────────────────────────────────────────────────────────

heading(doc, '4. Back-end REST API', level=1)

para(doc, (
    'De back-end biedt 70+ REST-endpoints gegroepeerd in 8 route-modules. '
    'Alle endpoints (behalve /api/auth/login en /health) zijn beveiligd met JWT-middleware.'
))

heading(doc, '4.1 Route-overzicht', level=2)

two_col_table(doc,
    rows=[
        ('/api/auth/*',             '2 endpoints — login, huidige gebruiker opvragen'),
        ('/api/notifications/*',    '22 endpoints — CRUD meldingen, beslissingen, bulk-beslissing, commentaren, bijlagen, coördinaten, zone-detectie, product-detectie'),
        ('/api/tasks/*',            '20 endpoints — CRUD taken, workflow, commentaren, artikelen, HPD-sync, productielijn-status, gerelateerde taken'),
        ('/api/products/*',         '8 endpoints — CRUD producten, koppelen/ontkoppelen aan meldingen, detectie op geometrie'),
        ('/api/product-versions/*', '6 endpoints — versies aanmaken, publiceren, taken koppelen, incomplete taken doorschuiven'),
        ('/api/production-lines/*', '2 endpoints — lijst productielijnen'),
        ('/api/users/*',            '3 endpoints — lijst gebruikers, rechten'),
        ('/api/coverages/*',        '5 endpoints — KML-dekgebieden beheren'),
    ],
    header=['Route prefix', 'Omschrijving'],
    col_widths=[2.5, 4.3],
)

heading(doc, '4.2 Sleutelendpoints in detail', level=2)

heading(doc, 'Meldingen', level=3)
bullet(doc, 'POST /api/notifications — Melding aanmaken; zone-detectie wordt automatisch getriggerd.')
bullet(doc, 'POST /api/notifications/:id/decide — Beslissing nemen (Ja/Nee); bij "Ja" wordt automatisch een taak aangemaakt en producten gekoppeld.')
bullet(doc, 'POST /api/notifications/bulk-decide — Bulk-beslissing op meerdere meldingen tegelijk.')
bullet(doc, 'POST /api/notifications/:id/detect-zones — Manuele herdetectie van zones.')
bullet(doc, 'GET /api/notifications/:id/attachments — Bijlagen ophalen; bestanden staan op /uploads/.')

heading(doc, 'Taken', level=3)
bullet(doc, 'PUT /api/tasks/:id — Taak bijwerken (formulier incl. HPD-velden, artikelen, workflow).')
bullet(doc, 'PUT /api/tasks/:taskId/products/:productId — Productstatus per taak bijwerken (te_verwerken → voltooid).')
bullet(doc, 'POST /api/tasks/:id/translate — Artikeltekst vertalen via Google Cloud Translate.')
bullet(doc, 'GET /api/tasks/:id/hpd-project — CARIS HPD-projectinfo ophalen.')

heading(doc, 'Productversies', level=3)
bullet(doc, 'POST /api/product-versions/:id/publish — Versie publiceren; niet-voltooide taken worden automatisch doorgeschoven naar de volgende versie.')

heading(doc, '4.3 Beveiliging', level=2)

two_col_table(doc,
    rows=[
        ('JWT authenticatie',           '✅ Bearer token, 24h geldigheidsduur'),
        ('Wachtwoord-hashing',          '✅ bcryptjs'),
        ('SQL-injectie preventie',      '✅ Parameterized queries (pg)'),
        ('CORS',                        '✅ Geconfigureerd in Express'),
        ('Security headers',            '✅ Helmet middleware'),
        ('Role-based access control',   '✅ can_view / can_edit / can_publish per productielijn'),
        ('File upload validatie',       '✅ MIME-typecheck + max 10 MB'),
        ('Auto-logout bij 401',         '✅ Front-end interceptor'),
        ('Rate limiting',               '⚠️ Nog niet geïmplementeerd'),
        ('Input sanitization',          '⚠️ Deels via express-validator; verder uitbreiden'),
        ('CSRF-beveiliging',            '⚠️ Nog niet geïmplementeerd'),
    ],
    header=['Beveiligingsmaatregel', 'Status'],
    col_widths=[3.2, 3.6],
)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 5. Front-end ─────────────────────────────────────────────────────────────

heading(doc, '5. Front-end Applicatie', level=1)

heading(doc, '5.1 Paginastructuur', level=2)

two_col_table(doc,
    rows=[
        ('/login',                      'Inlogformulier met validatie; redirect naar dashboard na succesvolle login'),
        ('/ (Dashboard)',               'Statistieken: openstaande meldingen, actieve taken; overzichtswidgets per productielijn'),
        ('/notifications',              'Meldingenlijst (paginering, zoeken, filteren, sorteerbare kolommen, bulk-beslissing)'),
        ('/notifications/:id',          'Meldingdetail: kaart (14 WMS-lagen), bijlagen, commentaren, coördinaten, zones, gekoppelde producten'),
        ('/tasks',                      'Takenlijst per productielijn (multi-productielijn status, filtering, sortering)'),
        ('/tasks/:id',                  'Taakdetail: formulier, kaart, workflow, artikelen (WYSIWYG), HPD-sync, gerelateerde taken'),
        ('/products',                   'Productencatalogus (tabel- en kaartweergave, kleurcodering per type/gebruiksniveau)'),
        ('/product-versions',           'Productversies (status In Progress / Ready / Published, taken koppelen)'),
        ('/published-product-versions', 'Archief van gepubliceerde versies'),
    ],
    header=['Route', 'Inhoud'],
    col_widths=[2.8, 4.0],
)

heading(doc, '5.2 State Management', level=2)
bullet(doc, 'Zustand authStore: token, gebruikersobject, actieve productielijn-ID — persistent via localStorage.')
bullet(doc, 'TanStack React Query: server-state caching, auto-retry, background refetch voor alle API-aanroepen.')
bullet(doc, 'Productielijn-selector is verplicht; zonder selectie worden gefilterde endpoints geblokkeerd.')

heading(doc, '5.3 Herbruikbare componenten', level=2)

two_col_table(doc,
    rows=[
        ('Layout',            'Header, inklapbare sidebar, content-area; productielijn-switcher met waarschuwingsindicator'),
        ('CoordinateInput',   '7 coördinaatformaten: DD, DDM, DMS + 4 geprojecteerde CRS (via proj4)'),
        ('FileUpload',        'Drag & drop, MIME-validatie, max 10 MB; bestandslijst met download/verwijder'),
        ('TaskCharts',        'SVG-taartdiagram voor statusverdeling van taken'),
        ('QuillEditor',       'Rich-text editor voor commentaren'),
        ('SunEditor',         'WYSIWYG voor meldingstoelichting en artikelen'),
    ],
    header=['Component', 'Beschrijving'],
    col_widths=[2.0, 4.8],
)

heading(doc, '5.4 Kaartintegratie (Leaflet)', level=2)
bullet(doc, '14 WMS-lagen beschikbaar (zeekaarten, IENC, bathymetrie, enz.).')
bullet(doc, 'GeoJSON-visualisatie van meldinggeometrie en productomtrekken op kaart.')
bullet(doc, 'Leaflet Draw tools voor handmatige invoer/bewerking van geometrieën.')
bullet(doc, 'Shapefile-export van coördinaten via JSZip + @mapbox/shp-write.')

heading(doc, '5.5 Sorteerfunctionaliteit', level=2)
para(doc, (
    'De aangepaste useTableSort hook voorziet client-side sortering op alle tabellen '
    '(drie stadia: oplopend → aflopend → geen sortering). Ondersteunde datatypes: '
    'string (locale-aware), getal, datum, boolean. Null-waarden worden altijd onderaan geplaatst.'
))

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 6. Hoofdprocessen ───────────────────────────────────────────────────────

heading(doc, '6. Functionele Hoofdprocessen', level=1)

heading(doc, '6.1 Vier kernprocessen', level=2)

para(doc, 'CARTIS 2.0 organiseert het werk rond vier processen die elke productielijn doorloopt:')
bullet(doc, 'Meldingen — inkomende nautische informatie registreren en beoordelen.')
bullet(doc, 'Taken — werkitems aanmaken op basis van beslissingen bij meldingen.')
bullet(doc, 'Productversies — taken verwerken in versies van nautische producten.')
bullet(doc, 'Publicaties — het effectief publiceren van de producten.')

heading(doc, '6.2 Procesflow', level=2)
code_block(doc, 'Melding binnenkomt')
code_block(doc, '  → Automatische zone-detectie (point-in-polygon)')
code_block(doc, '  → Automatische product-detectie (PostGIS overlap)')
code_block(doc, '  → Beslissing per productielijn (Ja / Nee / –)')
code_block(doc, '      → Bij "Ja": Taak aanmaken (taaknummer YYNNN)')
code_block(doc, '          → Producten koppelen aan taak')
code_block(doc, '          → Productstatus bijhouden (te_verwerken → voltooid)')
code_block(doc, '          → CARIS HPD-project synchroniseren')
code_block(doc, '          → Artikelen opstellen (WYSIWYG, Google Translate)')
code_block(doc, '      → Taken verzamelen in Productversie')
code_block(doc, '          → Status: In Progress → Ready → Published')
code_block(doc, '          → Incomplete taken doorschuiven naar volgende versie')

heading(doc, '6.3 Productielijnen', level=2)

two_col_table(doc,
    rows=[
        ('ZK',        'Zeekaartproductie — elektronische en papieren nautische kaarten (ENC_*.kml)'),
        ('IENC',      'Inland ENC — binnenvaartkaarten (IENC*.kml)'),
        ('PILOT_ENC', 'Pilot ENC — gedetailleerde bathymetrische loodskaarten (Pilot-ENC*.kml)'),
        ('PUBL',      'Publicaties — Berichten aan Zeevarenden (BaZ), Lichtenlijst, Verbeterlijst, enz.'),
    ],
    header=['Code', 'Omschrijving'],
    col_widths=[1.5, 5.3],
)

heading(doc, '6.4 Meldingen – bronnen en invoerkanalen', level=2)

two_col_table(doc,
    rows=[
        ('REST-API',        'Automatische bevraging van nautische bericht-API\'s (MRCC, BASS, POAB, FLARIS NtS) — configuratie aanwezig, import in ontwikkeling'),
        ('Mailbox',         'Automatisch inlezen uit nautinfo@mow.vlaanderen.be — Nodemailer geconfigureerd, import in ontwikkeling'),
        ('Handmatige invoer','Formulier + drag-and-drop van mails; velden automatisch ingevuld uit mail-metadata; bijlagen meeopgeslagen'),
        ('BaZ1-berichten',  'Gepubliceerde BaZ1-artikels overgenomen als nieuwe melding voor het volgende jaar'),
        ('Pushberichten',   'Wrakkendatabank, POSEIDON peilingen — gepland'),
    ],
    header=['Kanaal', 'Status / Beschrijving'],
    col_widths=[1.8, 5.0],
)

heading(doc, '6.5 Taaknummergeneratie', level=2)
para(doc, (
    'Taaknummers zijn uniek over alle productielijnen heen. Structuur: '
    'laatste twee cijfers van het jaar + viercijferig volgnummer (bv. 250006). '
    'Bij "Ja"-beslissing wordt het taaknummer aangemaakt en het bijbehorende CARIS HPD-project '
    'bepaald. Indien een andere productielijn al een taaknummer heeft gekoppeld, wordt datzelfde '
    'nummer hergebruikt.'
))

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 7. Geïmplementeerde functionaliteiten ────────────────────────────────────

heading(doc, '7. Geïmplementeerde Functionaliteiten', level=1)

heading(doc, '7.1 Meldingen', level=2)
features_notif = [
    'Lijstweergave met filtering op status, bron en datum',
    'Zoeken in titel, code en inhoud',
    'Beslissingen per productielijn (enkelvoudig + bulk)',
    'Automatische taak-aanmaak bij "Ja"-beslissing',
    'Producten koppelen (manueel) en automatisch detecteren (PostGIS)',
    'Paginering',
    'Bijlagen uploaden (drag & drop, max 10 MB), beheren en downloaden',
    'Detailpagina met Leaflet-kaart (14 WMS-lagen)',
    'Inklapbare lijstitems met volledige details',
    'SunEditor WYSIWYG-editor voor toelichtingsveld',
    'Commentaarsysteem (toevoegen, bewerken, verwijderen)',
    'Extra coördinaten toevoegen (DD/DDM/DMS/geprojecteerd) + GML-export',
    'Automatische zone-detectie + manuele zone-beheer',
    'Zone-badges (blauw = automatisch, oranje = manueel)',
    'Opvolgingsvlaggen (MSI actief, opvolging nodig, extra info nodig)',
]
for f in features_notif:
    bullet(doc, f'✅  {f}')

heading(doc, '7.2 Taken', level=2)
features_tasks = [
    'Lijstweergave per productielijn met multi-productielijn statusweergave',
    'Filteren op status; zoeken op taaknummer en titel',
    'Taaknummergeneratie (jaar + volgnummer)',
    'BaZ-nummer tracking',
    'MSI actief indicator',
    'Opvolgingsvlaggen',
    'Productstatus per taak bijhouden (te_verwerken → voltooid)',
    'Gerelateerde taken koppelen',
    'Koppeling met meldingen',
    'Commentaarsysteem',
    'Workflow-statussen',
    'Productielijn-specifieke statussen',
    'Taakdetailformulier met Leaflet-kaart (14 WMS-lagen)',
    'Artikelenbeheer (WYSIWYG) met Google Translate-vertaling',
    'CARIS HPD-projectkoppeling en synchronisatie',
]
for f in features_tasks:
    bullet(doc, f'✅  {f}')

heading(doc, '7.3 Producten', level=2)
features_prods = [
    'Lijstweergave per productielijn',
    'Producttypes en actief/inactief-status',
    'GeoJSON-geometrie (geografische omtrek)',
    'Tabel- en kaartweergave toggle',
    'Kleurcodering per producttype en gebruiksniveau',
    'KML-import (automatische productielijntoewijzing op basis van bestandsnaam)',
    'Koppelen/ontkoppelen aan meldingen (incl. PostGIS-overlap detectie)',
    'Sorteerbare kolommen',
]
for f in features_prods:
    bullet(doc, f'✅  {f}')

heading(doc, '7.4 Productversies', level=2)
features_pv = [
    'Versiebeheer per product',
    'Status tracking (In Progress / Ready / Published)',
    'Publicatiedatum registratie',
    'Taken koppelen aan versies',
    'Automatisch doorschuiven van onvolledige taken bij publicatie',
    'Archief van gepubliceerde versies',
    'Sorteerbare kolommen (versienummer, datum, status)',
]
for f in features_pv:
    bullet(doc, f'✅  {f}')

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 8. Openstaande werkpunten ───────────────────────────────────────────────

heading(doc, '8. Openstaande Werkpunten & Roadmap', level=1)

heading(doc, '8.1 Must-have (Prio 1) — Voltooid', level=2)
done_p1 = [
    'Bijlagebeheer met drag & drop',
    'Detailpagina\'s voor meldingen en taken',
    'Bewerkfunctionaliteit voor taken (uitgebreid detailformulier)',
    'WYSIWYG-editor voor toelichtingen (SunEditor)',
    'Geografisch overzicht (Leaflet-kaarten met 14 WMS-lagen)',
]
for d in done_p1:
    bullet(doc, f'✅  {d}')

heading(doc, '8.2 Should-have (Prio 2)', level=2)
todo_p2 = [
    ('✅', 'CARIS HPD-integratie — HPD-projectservice + endpoints aanwezig'),
    ('✅', 'Bulk-acties op meldingen (bulk-decide endpoint)'),
    ('✅', 'Exportfunctionaliteit (Shapefile via JSZip + shp-write)'),
    ('✅', 'BaZ1-artikelenbeheer + Google Translate-vertaling'),
    ('⚠️', 'Automatische mail-import (nautinfo mailbox) — Nodemailer geconfigureerd, import nog niet volledig'),
    ('⚠️', 'REST API-integraties (MRCC, BASS, POAB, FLARIS) — .env-configuratie aanwezig, endpoints in ontwikkeling'),
]
for status, desc in todo_p2:
    bullet(doc, f'{status}  {desc}')

heading(doc, '8.3 Nice-to-have (Prio 3)', level=2)
todo_p3 = [
    'Pushberichten (Wrakkendatabank, POSEIDON)',
    'Doorlooptijden-tracking (tijdregistratie/flow)',
    'E-mailnotificaties',
    'Gebruikersbeheer UI',
    'Geavanceerde filtering',
    'Rapportages en dashboards',
    'Dark mode',
]
for t in todo_p3:
    bullet(doc, f'🔲  {t}')

heading(doc, '8.4 Technische schuld', level=2)
tech_debt = [
    'Unit- en end-to-end-tests',
    'API-documentatie (Swagger/OpenAPI)',
    'Docker-containers',
    'CI/CD-pipeline',
    'Error monitoring',
    'Rate limiting',
    'Volledige input sanitization',
    'CSRF-beveiliging',
    'Performance-optimalisatie (paginering server-side voor grote datasets)',
]
for t in tech_debt:
    bullet(doc, f'🔲  {t}')

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 9. Externe integraties ──────────────────────────────────────────────────

heading(doc, '9. Externe Integraties & Connectoren', level=1)

two_col_table(doc,
    rows=[
        ('CARIS HPD',           'Synchronisatie van projectgegevens; HPDProjectService + dedicated endpoints. Status: geïmplementeerd.'),
        ('Google Cloud Translate', 'Automatische vertaling van BaZ-artikelen. Status: geïmplementeerd.'),
        ('Nautinfo mailbox',    'nautinfo@mow.vlaanderen.be — Nodemailer geconfigureerd; automatische import in ontwikkeling.'),
        ('MRCC Oostende (MSI)', 'Maritime Safety Information API; .env-configuratie aanwezig; endpoints in ontwikkeling.'),
        ('VTS Scheldt (BASS)',  'Bekendmaking aan de Scheepvaart Scheldegebied; in ontwikkeling.'),
        ('Port of Antwerp-Bruges (POAB)', 'Nautische berichten; in ontwikkeling.'),
        ('EuRIS (FLARIS NtS)',  'Inland waterway notices; in ontwikkeling.'),
        ('Wrakkendatabank',     'Pushberichten voor wrakaanpassingen; gepland.'),
        ('POSEIDON / CARIS BDB','Peilinformatie (GML/XML); gepland.'),
    ],
    header=['Systeem', 'Status & beschrijving'],
    col_widths=[2.2, 4.6],
)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 10. Installatie & Deployment ────────────────────────────────────────────

heading(doc, '10. Installatie & Opstarten', level=1)

heading(doc, '10.1 Vereisten', level=2)
bullet(doc, 'Node.js 20+')
bullet(doc, 'PostgreSQL 15+ met PostGIS extensie')
bullet(doc, 'Python 3 (optioneel, voor analysescripts)')
bullet(doc, 'Windows (scripts beschikbaar als .ps1 en .bat)')

heading(doc, '10.2 Installatiestappen', level=2)
steps = [
    ('1', 'Database aanmaken',  'psql -U postgres -c "CREATE DATABASE cartis;"'),
    ('2', 'Schema uitvoeren',   'psql -U postgres -d cartis -f backend/database/schema.sql'),
    ('3', 'PostGIS inschakelen','psql -U postgres -d cartis -f backend/database/enable-postgis.sql'),
    ('4', 'Dependencies',       'npm install (monorepo root installeert backend + frontend)'),
    ('5', 'Environment',        'Kopieer backend/.env.example → backend/.env en pas DB-credentials aan'),
    ('6', 'KML-import (opt.)',  '.\\import-products-kml.ps1'),
    ('7', 'Starten',            '.\\start-cartis.ps1  OF  npm run dev'),
]
table = doc.add_table(rows=len(steps) + 1, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Stap', 'Actie', 'Commando']):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, '1A568C')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for ri, (step, action, cmd) in enumerate(steps):
    row = table.rows[ri + 1]
    row.cells[0].text = step
    row.cells[1].text = action
    row.cells[2].text = cmd
    if ri % 2 == 1:
        for cell in row.cells:
            set_cell_bg(cell, 'EEF4FB')

heading(doc, '10.3 Toegang', level=2)
bullet(doc, 'Frontend: http://localhost:5173')
bullet(doc, 'Backend API: http://localhost:3000')
bullet(doc, 'Health check: http://localhost:3000/health')
bullet(doc, 'Standaard login: test@cartis.be / test123')

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# ─── 11. Samenvatting ────────────────────────────────────────────────────────

heading(doc, '11. Samenvatting', level=1)

para(doc, (
    'CARTIS 2.0 is een volledig functionele webapplicatie voor de administratieve opvolging van '
    'nautische meldingen bij Vlaamse Hydrografie. De applicatie dekt de vier kernprocessen '
    '(Meldingen, Taken, Productversies, Publicaties) voor vier productielijnen (ZK, IENC, PILOT_ENC, PUBL). '
    'De technische stack — Node.js/Express back-end met PostgreSQL/PostGIS en React/TypeScript front-end — '
    'biedt een solide basis voor verdere uitbreiding.'
))

doc.add_paragraph()

para(doc, 'Sterke punten van de huidige implementatie:', bold=True)
strengths = [
    'Volledig operationele monorepo met gescheiden back-end en front-end',
    'Robuust databaseschema (20+ tabellen) met audit trail en rol-gebaseerde rechten',
    '70+ REST-endpoints met JWT-beveiliging en parameterized queries',
    'Ruimtelijke functionaliteit via PostGIS (zone-detectie, product-overlap)',
    'Rijke front-end met interactieve Leaflet-kaarten, 14 WMS-lagen en coördinaatvarianten',
    'KML-import voor producten, shapefile-export voor coördinaten',
    'CARIS HPD-koppeling en Google Translate-integratie voor meertalige artikelen',
    'Uitgebreide documentatie (README, PROJECT_STATUS, DEVELOPMENT, ZONE_DETECTION, PRODUCTS_*)',
]
for s in strengths:
    bullet(doc, f'✅  {s}')

doc.add_paragraph()
para(doc, 'Aandachtspunten voor de volgende fase:', bold=True)
attentions = [
    'Implementeer rate limiting en volledige input sanitization (OWASP Prio)',
    'Automatische mail-import en REST API-integraties (MRCC, BASS, POAB, FLARIS) voltooien',
    'Unit- en integratietests toevoegen',
    'CI/CD-pipeline en Docker-containerisatie opzetten',
    'Swagger/OpenAPI-documentatie genereren',
    'Doorlooptijden-tracking en pushberichten (Wrakkendatabank, POSEIDON) plannen',
]
for a in attentions:
    bullet(doc, f'🔲  {a}')

doc.add_paragraph()
doc.add_paragraph()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('CARTIS 2.0 — Analyse & Technische Documentatie — April 2026 — Vlaamse Hydrografie')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# ─── Save ─────────────────────────────────────────────────────────────────────

output_path = r'd:\Programming\Webapps\Cartis_new\CARTIS2_Analyse_April2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
