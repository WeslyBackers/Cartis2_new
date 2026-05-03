"""
Generate APP_FUNCTIONAL_ANALYSIS.docx from APP_FUNCTIONAL_ANALYSIS.md
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = r"d:\Programming\Webapps\Cartis_new\APP_FUNCTIONAL_ANALYSIS.md"
DOCX_PATH = r"d:\Programming\Webapps\Cartis_new\APP_FUNCTIONAL_ANALYSIS.docx"


def set_paragraph_shading(paragraph, fill_hex: str):
    """Apply background shading to a paragraph (used for code-block style)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_horizontal_rule(doc: Document):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def apply_inline_formatting(paragraph, text: str):
    """Parse inline bold/italic/code markers and add runs."""
    # Split on **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif full.startswith("*"):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif full.startswith("`"):
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def build_document(md_text: str) -> Document:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ── Default body style ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    lines = md_text.splitlines()
    i = 0

    def add_title_page():
        p = doc.add_heading("CARTIS 2.0", 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph("Full Functional Analysis")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(16)
        sub.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        doc.add_paragraph("")
        doc.add_page_break()

    add_title_page()

    while i < len(lines):
        line = lines[i]

        # ── Heading 1 (# ...) ─────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Heading 2 (## ...) ───────────────────────────────────────────────
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Heading 3 (### ...) ──────────────────────────────────────────────
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # ── Horizontal rule (---) ─────────────────────────────────────────────
        if re.match(r"^-{3,}\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Bullet list (- ...) ───────────────────────────────────────────────
        if line.startswith("- "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            apply_inline_formatting(p, text)
            i += 1
            continue

        # ── Numbered list (N. ...) ─────────────────────────────────────────────
        m_num = re.match(r"^\d+\.\s+(.*)", line)
        if m_num:
            text = m_num.group(1)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            apply_inline_formatting(p, text)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        text = line.strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        apply_inline_formatting(p, text)
        i += 1

    return doc


def main():
    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    doc = build_document(md_text)
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
